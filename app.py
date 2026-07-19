import os
import re
import base64
import ast
import json
import math
import operator
import random
import threading
from io import BytesIO
from pathlib import Path
from zipfile import ZipFile

import gradio as gr
import pandas as pd
import requests
from smolagents import (
    LiteLLMModel,
    Tool,
    ToolCallingAgent,
    WikipediaSearchTool,
)


DEFAULT_API_URL = "https://agents-course-unit4-scoring.hf.space"
RESULT_COLUMNS = ["Task ID", "Question", "Submitted Answer"]
HTTP_TIMEOUT = 45
MAX_EXTRACTED_CHARS = 8_000
WEBPAGE_CONNECT_TIMEOUT = 8
WEBPAGE_READ_TIMEOUT = 20
MAX_WEBPAGE_CHARS = 3_500
DEFAULT_GEMINI_MODEL = "gemini-3.5-flash"
DEFAULT_GEMINI_REVIEW_MODEL = DEFAULT_GEMINI_MODEL
PRACTICALLY_UNLIMITED_STEPS = 2_147_483_647
GAIA_ROWS_API = "https://datasets-server.huggingface.co/rows"
TASK_FILE_CACHE = {}
ATTACHMENT_CACHE = {}
WEBPAGE_CACHE = {}
SEARCH_CACHE = {}
SEARCH_LOCK = threading.Lock()
OFFICIAL_QUESTIONS_CACHE = []
QUESTIONS_CACHE_LOCK = threading.Lock()
QUESTIONS_SOURCE = ""


def compact_error(exc: Exception) -> str:
    message = str(exc).strip()
    return message or repr(exc)


class ResearchStagnationError(RuntimeError):
    """Raised when tool usage is no longer producing useful new evidence."""


class RunEvidenceTracker:
    """Thread-safe evidence recorder with loop/stagnation detection."""

    def __init__(self):
        self.lock = threading.Lock()
        self.interrupt_callback = None
        self.max_tool_calls = self._env_int("GAIA_STAGNATION_TOOL_CALLS", 14)
        self.max_discovery_calls = self._env_int(
            "GAIA_STAGNATION_SEARCH_CALLS", 8
        )
        self.max_duplicate_observations = self._env_int(
            "GAIA_STAGNATION_DUPLICATES", 2
        )
        self.reset()

    @staticmethod
    def _env_int(name: str, default: int) -> int:
        try:
            return max(1, int(str(os.getenv(name) or default)))
        except (TypeError, ValueError):
            return default

    def reset(self, question: str = ""):
        with self.lock:
            self.question = str(question or "")
            self.evidence = []
            self.seen_observations = set()
            self.tool_calls = 0
            self.discovery_calls = 0
            self.duplicate_observations = 0
            self.low_value_streak = 0
            self.stagnated = False
            self.stagnation_reason = ""

    @staticmethod
    def _observation_signature(text: str) -> str:
        normalized = re.sub(r"\s+", " ", text.casefold())
        normalized = re.sub(r"\b\d+(?:\.\d+)?s\b", "", normalized)
        return normalized[:2_500]

    def record(self, tool_name: str, value) -> str:
        text = str(value or "").strip()
        callback = None
        with self.lock:
            self.tool_calls += 1
            if tool_name in {"web_search", "wikipedia_search"}:
                self.discovery_calls += 1

            if text:
                signature = self._observation_signature(text)
                duplicate = signature in self.seen_observations
                if duplicate:
                    self.duplicate_observations += 1
                else:
                    self.seen_observations.add(signature)
                self.evidence.append(
                    {
                        "tool": str(tool_name),
                        "text": text[:5_000],
                    }
                )

                lowered = text.casefold()
                low_value = duplicate or any(
                    cue in lowered
                    for cue in (
                        "duplicate search skipped",
                        "no search results found",
                        "no wikipedia page found",
                        "no html tables were found",
                        "error fetching",
                        "web search failed",
                        "rate limited",
                    )
                )
                self.low_value_streak = (
                    self.low_value_streak + 1 if low_value else 0
                )

            reason = ""
            if self.duplicate_observations >= self.max_duplicate_observations:
                reason = (
                    "resultados de ferramentas repetidos "
                    f"{self.duplicate_observations} vezes"
                )
            elif self.low_value_streak >= 3:
                reason = (
                    f"{self.low_value_streak} ferramentas consecutivas sem "
                    "evidência útil"
                )
            elif self.discovery_calls >= self.max_discovery_calls:
                reason = (
                    f"{self.discovery_calls} buscas de descoberta sem "
                    "resposta final"
                )
            elif self.tool_calls >= self.max_tool_calls:
                reason = (
                    f"{self.tool_calls} chamadas de ferramenta sem "
                    "resposta final"
                )

            if reason and not self.stagnated:
                self.stagnated = True
                self.stagnation_reason = reason
                callback = self.interrupt_callback

        if callable(callback):
            callback()
        return value

    def current(self, max_chars: int = 18_000) -> str:
        with self.lock:
            items = list(self.evidence)
            question = self.question
        if not items:
            return ""

        terms = {
            token.casefold()
            for token in re.findall(
                r"[A-Za-zÀ-ÿ0-9][A-Za-zÀ-ÿ0-9'._-]{2,}",
                question,
            )
            if token.casefold()
            not in {
                "about", "between", "from", "have", "how", "many",
                "question", "that", "the", "their", "this", "what",
                "when", "where", "which", "with",
            }
        }
        scored = []
        for index, item in enumerate(items):
            lowered = item["text"].casefold()
            hits = sum(1 for term in terms if term in lowered)
            useful = not any(
                cue in lowered
                for cue in (
                    "duplicate search skipped",
                    "no search results found",
                    "error fetching",
                    "rate limited",
                )
            )
            scored.append((hits * 5 + int(useful), index, item))

        selected = sorted(scored, reverse=True)[:10]
        selected.sort(key=lambda entry: entry[1])
        blocks = []
        used = 0
        for _, index, item in selected:
            block = f"[{index + 1}. {item['tool']}]\n{item['text']}"
            remaining = max_chars - used
            if remaining <= 100:
                break
            blocks.append(block[:remaining])
            used += len(blocks[-1]) + 2
        return "\n\n".join(blocks)


def _direct_gaia_questions() -> list[dict]:
    """Reproduces the official scoring Space filter through the dataset API."""
    token = os.getenv("HF_TOKEN")
    if not token:
        raise RuntimeError(
            "HF_TOKEN é necessário para carregar o dataset GAIA diretamente."
        )

    rows = []
    offset = 0
    total = None
    while total is None or offset < total:
        response = requests.get(
            GAIA_ROWS_API,
            headers={"Authorization": f"Bearer {token}"},
            params={
                "dataset": "gaia-benchmark/GAIA",
                "config": "2023_level1",
                "split": "validation",
                "offset": offset,
                "length": 100,
            },
            timeout=HTTP_TIMEOUT,
        )
        if response.status_code in {401, 403}:
            raise RuntimeError(
                "Acesso ao dataset GAIA negado. Aceite as condições em "
                "https://huggingface.co/datasets/gaia-benchmark/GAIA e "
                "confirme que HF_TOKEN possui permissão de leitura."
            )
        response.raise_for_status()
        payload = response.json()
        batch = payload.get("rows") or []
        total = int(payload.get("num_rows_total") or len(batch))
        if not batch:
            break
        rows.extend(
            wrapper.get("row", wrapper)
            for wrapper in batch
            if isinstance(wrapper, dict)
        )
        offset += len(batch)

    questions = []
    for item in rows:
        metadata = item.get("Annotator Metadata") or {}
        if isinstance(metadata, str):
            try:
                metadata = json.loads(metadata)
            except json.JSONDecodeError:
                metadata = {}
        try:
            number_of_tools = int(metadata.get("Number of tools"))
            number_of_steps = int(metadata.get("Number of steps"))
        except (TypeError, ValueError, AttributeError):
            continue
        if number_of_tools >= 3 or number_of_steps >= 6:
            continue

        task_id = str(item.get("task_id") or "").strip()
        question = str(item.get("Question") or "").strip()
        if not task_id or not question or item.get("Final answer") is None:
            continue
        public_item = {
            "task_id": task_id,
            "question": question,
            "Level": item.get("Level"),
            "file_name": item.get("file_name"),
        }
        questions.append(
            {
                key: value
                for key, value in public_item.items()
                if value is not None
            }
        )

    if not questions:
        raise RuntimeError(
            "O dataset GAIA foi acessado, mas nenhuma questão correspondeu "
            "ao filtro oficial (tools < 3 e steps < 6)."
        )
    return questions


def fetch_official_questions(force_refresh: bool = False) -> list[dict]:
    """Uses the scoring API first and reconstructs its 20-question set if down."""
    global OFFICIAL_QUESTIONS_CACHE
    global QUESTIONS_SOURCE

    with QUESTIONS_CACHE_LOCK:
        if OFFICIAL_QUESTIONS_CACHE and not force_refresh:
            return [dict(item) for item in OFFICIAL_QUESTIONS_CACHE]

        errors = []
        try:
            response = requests.get(
                f"{DEFAULT_API_URL}/questions",
                timeout=HTTP_TIMEOUT,
            )
            response.raise_for_status()
            questions = response.json()
            if not isinstance(questions, list) or not questions:
                raise ValueError("A API retornou uma lista vazia.")
            QUESTIONS_SOURCE = "API oficial do exercício"
        except Exception as exc:
            errors.append(f"scoring API: {compact_error(exc)}")
            try:
                questions = _direct_gaia_questions()
                QUESTIONS_SOURCE = (
                    "dataset GAIA direto (fallback da API indisponível)"
                )
                print(
                    "API de pontuação indisponível; questões reconstruídas "
                    "diretamente do dataset GAIA."
                )
            except Exception as fallback_exc:
                errors.append(
                    f"dataset fallback: {compact_error(fallback_exc)}"
                )
                raise RuntimeError(
                    "Não foi possível carregar as questões oficiais. "
                    + " | ".join(errors)
                ) from fallback_exc

        normalized = []
        for item in questions:
            if not isinstance(item, dict):
                continue
            task_id = str(item.get("task_id") or "").strip()
            question = str(
                item.get("question") or item.get("Question") or ""
            ).strip()
            if not task_id or not question:
                continue
            normalized_item = dict(item)
            normalized_item["task_id"] = task_id
            normalized_item["question"] = question
            normalized.append(normalized_item)
            TASK_FILE_CACHE[task_id] = str(
                normalized_item.get("file_name") or ""
            ).strip()

        if not normalized:
            raise RuntimeError("Nenhuma questão oficial válida foi carregada.")
        OFFICIAL_QUESTIONS_CACHE = normalized
        return [dict(item) for item in normalized]


def instrument_tool(tool: Tool, tracker: RunEvidenceTracker) -> Tool:
    """Records tool observations for deterministic review and failover."""
    original_forward = tool.forward

    def recorded_forward(*args, **kwargs):
        result = original_forward(*args, **kwargs)
        return tracker.record(tool.name, result)

    tool.forward = recorded_forward
    return tool


def concise_query_from_question(question: str, max_terms: int = 14) -> str:
    stopwords = {
        "about", "answer", "attached", "could", "from", "give", "have",
        "into", "just", "number", "please", "provide", "question", "same",
        "that", "their", "there", "these", "this", "under", "what", "when",
        "where", "which", "with", "would", "your",
    }
    terms = [
        token
        for token in re.findall(r"[A-Za-z0-9][A-Za-z0-9'._-]{2,}", question or "")
        if token.casefold() not in stopwords
    ]
    return " ".join(terms[:max_terms])


def focus_text(text: str, query: str, max_chars: int = 5_000) -> str:
    """Selects high-signal passages locally, before text reaches the LLM."""
    text = re.sub(r"\r\n?", "\n", str(text or ""))
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text).strip()
    if not text or len(text) <= max_chars:
        return text

    stopwords = {
        "about", "after", "again", "also", "article", "attached", "before",
        "could", "find", "from", "have", "into", "just", "mentions", "please",
        "provide", "question", "should", "that", "their", "there", "these",
        "this", "under", "what", "when", "where", "which", "with", "work",
        "would", "your",
    }
    terms = {
        token.lower()
        for token in re.findall(r"[A-Za-z0-9][A-Za-z0-9._-]{2,}", query or "")
        if token.lower() not in stopwords
    }
    blocks = [
        block.strip()
        for block in re.split(r"\n\s*\n", text)
        if block.strip()
    ]
    if not blocks:
        return text[:max_chars]

    scored = []
    for index, block in enumerate(blocks):
        lowered = block.lower()
        # Count distinct terms instead of every occurrence. This prevents a
        # huge navigation/table block containing the same year many times from
        # outranking a short passage containing the actual target terms.
        hits = sum(1 for term in terms if term in lowered)
        exact_bonus = 4 if query and query.lower() in lowered else 0
        signal_bonus = 2 if re.search(
            r"\b(acknowledg|award|grant|answer|result|total|page|pages)\b",
            lowered,
        ) else 0
        scored.append((hits * 3 + exact_bonus + signal_bonus, index))

    best_indices = [
        index
        for score, index in sorted(scored, reverse=True)
        if score > 0
    ][:10]
    if not best_indices:
        return (
            text[: max_chars * 2 // 3]
            + "\n\n[content omitted]\n\n"
            + text[-max_chars // 3 :]
        )

    selected = {}
    used = 0
    for best_index in best_indices:
        # Reserve space for the matching block before optional neighbors.
        for index in (best_index, best_index - 1, best_index + 1):
            if (
                index in selected
                or not 0 <= index < len(blocks)
                or used >= max_chars
            ):
                continue
            block = blocks[index]
            if index != best_index and len(block) > 1_000:
                block = block[:1_000]
            remaining = max_chars - used
            if remaining < 100:
                break
            if len(block) > remaining and terms:
                lowered_block = block.lower()
                positions = [
                    lowered_block.find(term)
                    for term in terms
                    if lowered_block.find(term) >= 0
                ]
                if positions:
                    center = min(positions)
                    start = max(0, center - remaining // 3)
                    block = block[start : start + remaining]
            selected[index] = block[:remaining]
            used += len(selected[index]) + 2
    return "\n\n".join(selected[index] for index in sorted(selected)).strip()


def clean_filename(value: str) -> str:
    value = Path(value).name
    return re.sub(r"[^A-Za-z0-9._-]+", "_", value) or "attachment"


def filename_from_response(response: requests.Response, task_id: str) -> str:
    disposition = response.headers.get("content-disposition", "")
    utf8_match = re.search(
        r"filename\*=UTF-8''([^;]+)", disposition, flags=re.IGNORECASE
    )
    plain_match = re.search(
        r'filename="?([^";]+)"?', disposition, flags=re.IGNORECASE
    )
    if utf8_match:
        from urllib.parse import unquote

        return clean_filename(unquote(utf8_match.group(1)))
    if plain_match:
        return clean_filename(plain_match.group(1))

    content_type = response.headers.get("content-type", "").lower()
    extensions = {
        "application/pdf": ".pdf",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document": ".docx",
        "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet": ".xlsx",
        "text/csv": ".csv",
        "text/plain": ".txt",
        "application/json": ".json",
        "image/png": ".png",
        "image/jpeg": ".jpg",
    }
    extension = next(
        (ext for mime, ext in extensions.items() if mime in content_type), ""
    )
    return f"{clean_filename(task_id)}{extension}"


def download_gaia_attachment(task_id: str) -> tuple[bytes, str]:
    """Baixa um anexo pela API do curso, com fallback para o dataset oficial."""
    task_id = str(task_id).strip()
    if task_id in ATTACHMENT_CACHE:
        return ATTACHMENT_CACHE[task_id]

    course_url = f"{DEFAULT_API_URL}/files/{task_id}"
    try:
        response = requests.get(course_url, timeout=HTTP_TIMEOUT)
        if response.ok:
            result = (
                response.content,
                filename_from_response(response, task_id),
            )
            ATTACHMENT_CACHE[task_id] = result
            return result
        course_error = (
            f"course file endpoint returned HTTP {response.status_code}"
        )
    except Exception as exc:
        course_error = f"course file endpoint failed: {compact_error(exc)}"
    questions = fetch_official_questions()
    item = next(
        (
            question
            for question in questions
            if str(question.get("task_id")) == task_id
        ),
        None,
    )
    filename = str((item or {}).get("file_name") or "").strip()
    if not filename:
        raise FileNotFoundError(f"No attachment exists for task {task_id}.")

    token = os.getenv("HF_TOKEN")
    if not token:
        raise RuntimeError(
            f"{course_error}; HF_TOKEN is required "
            "for the official GAIA dataset fallback."
        )

    from huggingface_hub import hf_hub_download

    errors = []
    for dataset_path in (
        f"2023/validation/{filename}",
        filename,
    ):
        try:
            local_path = hf_hub_download(
                repo_id="gaia-benchmark/GAIA",
                repo_type="dataset",
                filename=dataset_path,
                token=token,
            )
            result = (Path(local_path).read_bytes(), filename)
            ATTACHMENT_CACHE[task_id] = result
            return result
        except Exception as exc:
            errors.append(str(exc))

    raise RuntimeError(
        "Could not download the attachment from the course API or official "
        "GAIA dataset. Accept the dataset access conditions at "
        "https://huggingface.co/datasets/gaia-benchmark/GAIA and ensure "
        f"HF_TOKEN has read access. Details: {' | '.join(errors)}"
    )


def get_task_file_name(task_id: str) -> str:
    """Retorna o nome oficial do anexo ou uma string vazia."""
    task_id = str(task_id or "").strip()
    if not task_id:
        return ""
    if task_id in TASK_FILE_CACHE:
        return TASK_FILE_CACHE[task_id]

    try:
        for item in fetch_official_questions():
            item_task_id = str(item.get("task_id", "")).strip()
            TASK_FILE_CACHE[item_task_id] = str(
                item.get("file_name") or ""
            ).strip()
    except Exception as exc:
        print(f"Could not load task attachment metadata: {exc}")
        return ""

    return TASK_FILE_CACHE.get(task_id, "")


def extract_attachment_text(
    data: bytes, filename: str, query: str = ""
) -> str:
    """Extrai conteúdo legível dos formatos mais comuns do GAIA."""
    suffix = Path(filename).suffix.lower()

    if suffix == ".pdf":
        from pypdf import PdfReader

        reader = PdfReader(BytesIO(data))
        text = "\n\n".join(page.extract_text() or "" for page in reader.pages)
    elif suffix == ".docx":
        from docx import Document

        document = Document(BytesIO(data))
        paragraphs = [paragraph.text for paragraph in document.paragraphs]
        for table in document.tables:
            for row in table.rows:
                paragraphs.append(" | ".join(cell.text for cell in row.cells))
        text = "\n".join(paragraphs)
    elif suffix in {".xlsx", ".xlsm"}:
        from openpyxl import load_workbook

        workbook = load_workbook(BytesIO(data), read_only=True, data_only=True)
        lines = []
        for sheet in workbook.worksheets:
            lines.append(f"--- Sheet: {sheet.title} ---")
            for row in sheet.iter_rows(values_only=True):
                if any(value is not None for value in row):
                    lines.append(
                        " | ".join(
                            "" if value is None else str(value) for value in row
                        )
                    )
        text = "\n".join(lines)
    elif suffix in {".csv", ".tsv"}:
        separator = "\t" if suffix == ".tsv" else ","
        dataframe = pd.read_csv(
            BytesIO(data), sep=separator, encoding_errors="replace"
        )
        text = dataframe.to_csv(index=False)
    elif suffix in {
        ".txt",
        ".md",
        ".json",
        ".html",
        ".htm",
        ".xml",
        ".py",
    }:
        text = data.decode("utf-8", errors="replace")
        if suffix in {".html", ".htm"}:
            from bs4 import BeautifulSoup

            text = BeautifulSoup(text, "html.parser").get_text("\n")
    elif suffix in {".mp3", ".wav", ".flac", ".m4a", ".ogg"}:
        text = (
            "Audio attachment detected. Use transcribe_gaia_audio with the "
            "task_id instead of inspect_gaia_attachment."
        )
    elif suffix == ".zip":
        with ZipFile(BytesIO(data)) as archive:
            text = "Files inside ZIP:\n" + "\n".join(archive.namelist())
    elif suffix in {".png", ".jpg", ".jpeg", ".webp", ".gif"}:
        from PIL import Image

        image = Image.open(BytesIO(data))
        text = (
            f"Image attachment: {filename}; format={image.format}; "
            f"size={image.width}x{image.height}. "
            "This text-only model cannot inspect image pixels reliably."
        )
    else:
        preview = data[:2_000].decode("utf-8", errors="replace")
        text = (
            f"Unsupported attachment format {suffix or '(unknown)'}. "
            f"Decoded preview:\n{preview}"
        )

    text = text.strip()
    if not text:
        return f"The attachment {filename} was downloaded but contained no extractable text."
    tabular = suffix in {".xlsx", ".xlsm", ".csv", ".tsv"}
    if query and not tabular:
        text = focus_text(text, query, MAX_EXTRACTED_CHARS)
    elif len(text) > MAX_EXTRACTED_CHARS:
        head_size = MAX_EXTRACTED_CHARS * 2 // 3
        text = (
            text[:head_size]
            + "\n\n[attachment middle omitted]\n\n"
            + text[-(MAX_EXTRACTED_CHARS - head_size) :]
        )
    return text


class ConciseWebSearchTool(Tool):
    name = "web_search"
    description = (
        "Searches the public web without an API key. Returns at most five "
        "compact results. Use one precise query, then open the best source. "
        "Identical repeated searches are skipped automatically."
    )
    inputs = {
        "query": {
            "type": "string",
            "description": "One concise web search query.",
        }
    }
    output_type = "string"

    def forward(self, query: str) -> str:
        from ddgs import DDGS

        query = " ".join(str(query or "").split())
        if not query:
            return "Search query is empty."
        cache_key = query.casefold()

        # The lock also prevents parallel duplicate searches from both sending
        # the same network request and duplicating a large observation.
        with SEARCH_LOCK:
            if cache_key in SEARCH_CACHE:
                return (
                    "Duplicate search skipped; the identical results are "
                    "already present in this run. Open one of those URLs."
                )
            try:
                raw_results = list(DDGS().text(query, max_results=4))
                lines = []
                for index, result in enumerate(raw_results[:4], start=1):
                    title = " ".join(str(result.get("title") or "").split())
                    url = str(result.get("href") or result.get("url") or "").strip()
                    snippet = " ".join(str(result.get("body") or "").split())
                    lines.append(
                        f"{index}. {title[:180]}\n"
                        f"URL: {url}\n"
                        f"Snippet: {snippet[:240]}"
                    )
                output = "\n\n".join(lines) or "No search results found."
                SEARCH_CACHE[cache_key] = output
                return output
            except Exception as exc:
                return f"Web search failed: {compact_error(exc)}"


class WikipediaRevisionTool(Tool):
    name = "wikipedia_revision"
    description = (
        "Reads the last English Wikipedia revision on or before December 31 "
        "of a requested year, including raw table content that ordinary "
        "Wikipedia summaries may omit. Use it whenever a task specifies a "
        "Wikipedia version, snapshot, revision, or year."
    )
    inputs = {
        "title": {
            "type": "string",
            "description": "Exact or likely English Wikipedia page title.",
        },
        "year": {
            "type": "integer",
            "description": "Snapshot year, for example 2022.",
        },
        "query": {
            "type": "string",
            "description": "Short terms for the target section, row, or fact.",
        },
    }
    output_type = "string"

    def forward(self, title: str, year: int, query: str) -> str:
        title = " ".join(str(title or "").split())
        query = " ".join(str(query or "").split())
        try:
            year = int(year)
        except (TypeError, ValueError):
            return "Invalid Wikipedia revision year."
        if not title or year < 2001 or year > 2100:
            return "Provide a page title and a valid Wikipedia revision year."

        try:
            response = requests.get(
                "https://en.wikipedia.org/w/api.php",
                params={
                    "action": "query",
                    "prop": "revisions",
                    "titles": title,
                    "redirects": 1,
                    "rvstart": f"{year}-12-31T23:59:59Z",
                    "rvdir": "older",
                    "rvlimit": 1,
                    "rvprop": "ids|timestamp|content",
                    "rvslots": "main",
                    "format": "json",
                    "formatversion": 2,
                },
                headers={
                    "User-Agent": (
                        "GAIA-Course-Agent/1.0 "
                        "(educational project; Wikipedia revision reader)"
                    )
                },
                timeout=(WEBPAGE_CONNECT_TIMEOUT, WEBPAGE_READ_TIMEOUT),
            )
            response.raise_for_status()
            pages = response.json().get("query", {}).get("pages", [])
            page = pages[0] if pages else {}
            revisions = page.get("revisions") or []
            if page.get("missing") or not revisions:
                return (
                    f"No English Wikipedia revision was found for '{title}' "
                    f"on or before {year}-12-31."
                )
            revision = revisions[0]
            content = (
                revision.get("slots", {}).get("main", {}).get("content")
                or revision.get("*")
                or ""
            )
            if not content:
                return "The Wikipedia revision contained no readable wikitext."
            focused = focus_text(content, query, max_chars=7_000)
            return (
                f"Wikipedia page: {page.get('title', title)}\n"
                f"Revision ID: {revision.get('revid', 'unknown')}\n"
                f"Revision timestamp: {revision.get('timestamp', 'unknown')}\n"
                f"Snapshot rule: last revision on or before {year}-12-31\n\n"
                f"{focused}"
            )
        except Exception as exc:
            return (
                "Could not read the historical Wikipedia revision: "
                f"{compact_error(exc)}"
            )


class WikipediaFeaturedArticlesTool(Tool):
    name = "wikipedia_featured_articles"
    description = (
        "Reads the official English Wikipedia featured-article promotion "
        "table for one month and year, returning article titles and "
        "nominators. It can filter the rows by subject using each article's "
        "intro and categories. Use it instead of web search for questions "
        "about Featured Articles promoted in a specified month."
    )
    inputs = {
        "year": {
            "type": "integer",
            "description": "Promotion year, for example 2016.",
        },
        "month": {
            "type": "string",
            "description": "English month name, for example November.",
        },
        "subject": {
            "type": "string",
            "description": (
                "Optional subject to identify, for example dinosaur. Use an "
                "empty string to return every promotion in the month."
            ),
        },
    }
    output_type = "string"

    MONTHS = {
        name.casefold(): name
        for name in (
            "January", "February", "March", "April", "May", "June",
            "July", "August", "September", "October", "November", "December",
        )
    }

    @staticmethod
    def _wiki_link_value(cell: str, user: bool = False) -> str:
        if user:
            matches = re.findall(
                r"\[\[\s*(?:User|User talk)\s*:\s*([^|\]]+)"
                r"(?:\|([^\]]+))?\]\]",
                cell,
                flags=re.I,
            )
            values = [
                re.sub(r"<[^>]+>", "", display or target).strip()
                for target, display in matches
                if (display or target).strip()
            ]
            return " & ".join(dict.fromkeys(values))

        match = re.search(
            r"\[\[\s*([^|\]#]+)(?:#[^|\]]*)?(?:\|([^\]]+))?\]\]",
            cell,
        )
        if not match:
            return re.sub(r"<[^>]+>", "", cell).strip(" |")
        return re.sub(
            r"<[^>]+>",
            "",
            match.group(1),
        ).strip()

    @staticmethod
    def _article_metadata(titles: list[str]) -> dict[str, str]:
        metadata = {}
        for start in range(0, len(titles), 20):
            response = requests.get(
                "https://en.wikipedia.org/w/api.php",
                params={
                    "action": "query",
                    "prop": "extracts|categories",
                    "titles": "|".join(titles[start:start + 20]),
                    "redirects": 1,
                    "exintro": 1,
                    "explaintext": 1,
                    "exsentences": 2,
                    "cllimit": "max",
                    "format": "json",
                    "formatversion": 2,
                },
                headers={
                    "User-Agent": (
                        "GAIA-Course-Agent/1.0 "
                        "(educational project; featured article reader)"
                    )
                },
                timeout=(WEBPAGE_CONNECT_TIMEOUT, WEBPAGE_READ_TIMEOUT),
            )
            response.raise_for_status()
            for page in response.json().get("query", {}).get("pages", []):
                categories = " ".join(
                    category.get("title", "")
                    for category in page.get("categories", [])
                )
                metadata[str(page.get("title") or "")] = (
                    f"{page.get('extract') or ''} {categories}"
                ).casefold()
        return metadata

    def forward(self, year: int, month: str, subject: str = "") -> str:
        try:
            year = int(year)
        except (TypeError, ValueError):
            return "Invalid featured-article promotion year."
        normalized_month = self.MONTHS.get(
            " ".join(str(month or "").split()).casefold()
        )
        subject = " ".join(str(subject or "").split()).strip()
        if year < 2004 or year > 2100 or not normalized_month:
            return "Provide a valid year and an English month name."

        page_title = f"Wikipedia:Featured articles promoted in {year}"
        try:
            response = requests.get(
                "https://en.wikipedia.org/w/api.php",
                params={
                    "action": "parse",
                    "page": page_title,
                    "prop": "wikitext",
                    "format": "json",
                    "formatversion": 2,
                },
                headers={
                    "User-Agent": (
                        "GAIA-Course-Agent/1.0 "
                        "(educational project; featured article reader)"
                    )
                },
                timeout=(WEBPAGE_CONNECT_TIMEOUT, WEBPAGE_READ_TIMEOUT),
            )
            response.raise_for_status()
            payload = response.json()
            if "error" in payload:
                return (
                    "Wikipedia did not return the requested promotion page: "
                    f"{payload['error'].get('info', 'unknown error')}"
                )
            wikitext = payload.get("parse", {}).get("wikitext", "")
            section_match = re.search(
                rf"==\s*Promoted in {re.escape(normalized_month)} {year}\s*=="
                rf"(.*?)(?=\n==\s*Promoted in|\Z)",
                wikitext,
                flags=re.I | re.S,
            )
            if not section_match:
                return (
                    f"No promotion section was found for "
                    f"{normalized_month} {year}."
                )

            rows = []
            for row in re.split(r"\n\|-", section_match.group(1)):
                if row.count("||") < 3 or "Article!!" in row:
                    continue
                article = self._wiki_link_value(row)
                nominator = self._wiki_link_value(row, user=True)
                if article and nominator:
                    rows.append(
                        {
                            "article": article,
                            "nominator": nominator,
                        }
                    )
            if not rows:
                return (
                    f"No promotion rows could be parsed for "
                    f"{normalized_month} {year}."
                )

            selected = rows
            if subject:
                metadata = self._article_metadata(
                    [row["article"] for row in rows]
                )
                subject_terms = [
                    term.casefold()
                    for term in re.findall(r"[A-Za-z0-9-]{3,}", subject)
                ]
                selected = [
                    row
                    for row in rows
                    if all(
                        term in (
                            row["article"].casefold()
                            + " "
                            + metadata.get(row["article"], "")
                        )
                        for term in subject_terms
                    )
                ]

            heading = (
                f"English Wikipedia Featured Articles promoted in "
                f"{normalized_month} {year}.\n"
            )
            if subject:
                heading += (
                    f"Subject filter: {subject}. "
                    f"Matched {len(selected)} of {len(rows)} rows.\n"
                )
            if not selected:
                compact_rows = "; ".join(
                    f"{row['article']} — {row['nominator']}" for row in rows
                )
                return (
                    heading
                    + "No subject match was confirmed from article intros or "
                    "categories. Unfiltered rows:\n"
                    + compact_rows
                )
            return heading + "\n".join(
                f"- Article: {row['article']} | Nominator: {row['nominator']}"
                for row in selected
            )
        except Exception as exc:
            return (
                "Could not read the Wikipedia featured-article table: "
                f"{compact_error(exc)}"
            )


class OpenWebPageTool(Tool):
    name = "visit_webpage"
    description = (
        "Opens and reads one exact HTTP/HTTPS page. Use it after web_search to "
        "verify article text, tables, archives, papers, and linked sources. "
        "It returns concise main-page content while preserving links near the "
        "end, and retries blocked HTML through a text mirror. It is not a "
        "binary-file downloader."
    )
    inputs = {
        "url": {
            "type": "string",
            "description": "The complete HTTP or HTTPS URL to open.",
        },
        "query": {
            "type": "string",
            "description": (
                "Short terms describing the exact fact to find on the page. "
                "Do not repeat the full task."
            ),
        },
    }
    output_type = "string"

    def forward(self, url: str, query: str) -> str:
        from markdownify import markdownify

        url = str(url or "").strip()
        query = str(query or "").strip()
        if not re.match(r"^https?://", url, flags=re.I):
            return "Invalid URL: visit_webpage requires a full HTTP/HTTPS URL."
        if url in WEBPAGE_CACHE:
            return focus_text(WEBPAGE_CACHE[url], query, MAX_WEBPAGE_CHARS)

        headers = {
            "User-Agent": (
                "Mozilla/5.0 (compatible; GAIA-Course-Agent/1.0; "
                "+https://huggingface.co/learn/agents-course)"
            )
        }
        errors = []
        targets = [url]
        if "r.jina.ai/http" not in url:
            without_scheme = re.sub(r"^https?://", "", url)
            targets.append(f"https://r.jina.ai/http://{without_scheme}")

        for target in targets:
            try:
                response = requests.get(
                    target,
                    headers=headers,
                    timeout=(WEBPAGE_CONNECT_TIMEOUT, WEBPAGE_READ_TIMEOUT),
                    allow_redirects=True,
                )
                response.raise_for_status()
                content_type = response.headers.get("content-type", "").lower()
                if (
                    "text/" not in content_type
                    and "json" not in content_type
                    and "xml" not in content_type
                ):
                    return (
                        f"The URL returned binary content ({content_type}). "
                        "Use inspect_gaia_attachment for official GAIA files."
                    )

                text = response.text
                if "html" in content_type:
                    from bs4 import BeautifulSoup
                    from trafilatura import extract
                    from urllib.parse import urljoin

                    soup = BeautifulSoup(text, "html.parser")
                    for element in soup.select(
                        "script, style, noscript, nav, header, footer, form, "
                        "aside, iframe"
                    ):
                        element.decompose()
                    main_content = (
                        soup.find("article")
                        or soup.find("main")
                        or soup.body
                        or soup
                    )
                    extracted = extract(
                        text,
                        url=target,
                        output_format="markdown",
                        include_comments=False,
                        include_tables=True,
                        include_links=True,
                        favor_recall=True,
                    )
                    text = (
                        extracted
                        if extracted and len(extracted.strip()) >= 200
                        else markdownify(str(main_content))
                    )

                    # Article extractors sometimes classify bibliography links
                    # as navigation. Append a compact link index so the agent
                    # can still open papers and primary sources cited at the end.
                    source_links = []
                    seen_links = set()
                    for anchor in main_content.find_all("a", href=True):
                        label = " ".join(anchor.get_text(" ", strip=True).split())
                        absolute_url = urljoin(target, anchor["href"])
                        if (
                            label
                            and absolute_url.startswith(("http://", "https://"))
                            and absolute_url not in seen_links
                        ):
                            seen_links.add(absolute_url)
                            source_links.append(
                                f"- [{label[:160]}]({absolute_url})"
                            )
                    if source_links:
                        text += (
                            "\n\nSource links found on the page:\n"
                            + "\n".join(source_links[-12:])
                        )
                text = re.sub(r"\n{3,}", "\n\n", text).strip()
                if text:
                    WEBPAGE_CACHE[url] = text[:50_000]
                    return focus_text(text, query, MAX_WEBPAGE_CHARS)
                return "The page was retrieved but contained no text."
            except Exception as exc:
                errors.append(f"{target}: {compact_error(exc)}")

        return "Error fetching the webpage: " + " | ".join(errors)


class ReadDocumentUrlTool(Tool):
    name = "read_document_url"
    description = (
        "Downloads and reads a PDF, DOCX, CSV, or text document from an exact "
        "public URL, then returns only passages relevant to the supplied query. "
        "Use it for linked papers and reports; use visit_webpage for HTML."
    )
    inputs = {
        "url": {
            "type": "string",
            "description": "Direct public URL of the document.",
        },
        "query": {
            "type": "string",
            "description": "Short terms for the exact fact to find.",
        },
    }
    output_type = "string"

    def forward(self, url: str, query: str) -> str:
        url = str(url or "").strip()
        query = str(query or "").strip()
        if not re.match(r"^https?://", url, flags=re.I):
            return "Invalid document URL."

        try:
            response = requests.get(
                url,
                headers={"User-Agent": "GAIA-Course-Agent/1.0"},
                timeout=(WEBPAGE_CONNECT_TIMEOUT, HTTP_TIMEOUT),
                allow_redirects=True,
            )
            response.raise_for_status()
            if len(response.content) > 25 * 1024 * 1024:
                return "Document exceeds the 25 MB safety limit."

            content_type = response.headers.get("content-type", "").lower()
            if "text/html" in content_type:
                return (
                    "This URL returned HTML. Use visit_webpage with the same "
                    "URL and a short query."
                )
            filename = filename_from_response(response, "web_document")
            if not Path(filename).suffix:
                from urllib.parse import urlparse

                filename = Path(urlparse(response.url).path).name or filename
            extracted = extract_attachment_text(
                response.content, filename, query=query
            )
            return f"Document: {filename}\n\n{extracted}"
        except Exception as exc:
            return f"Could not read document URL: {compact_error(exc)}"


class QueryWebTableTool(Tool):
    name = "query_web_table"
    description = (
        "Reads an HTML table and computes its minimum or maximum locally, "
        "without sending the full table to the model. Use it for rankings, "
        "counts, standings, statistics, and 'least/most' questions after a "
        "table URL is known. Column selectors may be 'first', 'last', a "
        "zero-based number, or visible header text. Ties are alphabetized."
    )
    inputs = {
        "url": {
            "type": "string",
            "description": "Exact HTTP/HTTPS page containing the table.",
        },
        "operation": {
            "type": "string",
            "description": "One of: describe, min, or max.",
        },
        "label_column": {
            "type": "string",
            "description": "Label column: usually first, or its header text.",
        },
        "value_column": {
            "type": "string",
            "description": "Numeric column: last, a zero-based index, or header text.",
        },
    }
    output_type = "string"

    @staticmethod
    def _number(value: str):
        cleaned = str(value or "").strip().replace(",", "")
        cleaned = re.sub(r"[%+$]", "", cleaned)
        if not re.fullmatch(r"-?\d+(?:\.\d+)?", cleaned):
            return None
        number = float(cleaned)
        return int(number) if number.is_integer() else number

    @staticmethod
    def _column_index(selector: str, headers: list[str], default: int) -> int:
        selector = str(selector or "").strip().casefold()
        if selector == "first":
            return 0
        if selector == "last":
            return -1
        if re.fullmatch(r"-?\d+", selector):
            return int(selector)
        matches = [
            index
            for index, header in enumerate(headers)
            if selector and selector in header.casefold()
        ]
        return matches[-1] if matches else default

    def forward(
        self,
        url: str,
        operation: str,
        label_column: str,
        value_column: str,
    ) -> str:
        from bs4 import BeautifulSoup

        url = str(url or "").strip()
        operation = str(operation or "").strip().lower()
        if not re.match(r"^https?://", url, flags=re.I):
            return "Invalid URL."
        if operation not in {"describe", "min", "max"}:
            return "Invalid operation. Use describe, min, or max."

        try:
            response = requests.get(
                url,
                headers={"User-Agent": "GAIA-Course-Agent/1.0"},
                timeout=(WEBPAGE_CONNECT_TIMEOUT, WEBPAGE_READ_TIMEOUT),
            )
            response.raise_for_status()
            soup = BeautifulSoup(response.text, "html.parser")
            tables = soup.find_all("table")
            if not tables:
                return "No HTML tables were found on this page."

            table = max(tables, key=lambda item: len(item.find_all("tr")))
            parsed_rows = []
            header_rows = []
            for row in table.find_all("tr"):
                cells = row.find_all(["th", "td"])
                values = [
                    " ".join(cell.get_text(" ", strip=True).split())
                    for cell in cells
                ]
                if not values:
                    continue
                if row.find("th"):
                    header_rows.append(values)
                parsed_rows.append(values)

            if not parsed_rows:
                return "The largest table contained no usable rows."
            headers = max(header_rows, key=len) if header_rows else parsed_rows[0]
            if operation == "describe":
                return (
                    f"Rows: {len(parsed_rows)}; approximate columns: "
                    f"{len(max(parsed_rows, key=len))}; headers: "
                    + " | ".join(headers[:80])
                    + ". Use first/last when multi-row headers are ambiguous."
                )[:3_000]

            label_index = self._column_index(label_column, headers, 0)
            value_index = self._column_index(value_column, headers, -1)
            candidates = []
            for row in parsed_rows:
                try:
                    label = row[label_index].strip()
                    value_text = row[value_index].strip()
                except IndexError:
                    continue
                value = self._number(value_text)
                if value is None or not label or self._number(label) is not None:
                    continue
                candidates.append((value, label, value_text))

            if not candidates:
                return (
                    "No numeric rows matched those columns. Call describe and "
                    "retry with first/last or a numeric column index."
                )
            candidates.sort(
                key=lambda item: (
                    item[0] if operation == "min" else -item[0],
                    item[1].casefold(),
                )
            )
            best_value = candidates[0][0]
            tied = [item for item in candidates if item[0] == best_value]
            preview = candidates[:10]
            return (
                f"{operation.upper()} result (ties alphabetized): "
                f"{tied[0][1]} = {tied[0][2]}\n"
                "Top rows: "
                + "; ".join(
                    f"{label}={value_text}"
                    for _, label, value_text in preview
                )
            )
        except Exception as exc:
            return f"Could not query web table: {compact_error(exc)}"


class MlbStatsTool(Tool):
    name = "mlb_stats"
    description = (
        "Queries the official MLB Stats API without an API key and sorts a "
        "team's season hitting statistics locally. Use it for historical or "
        "current MLB questions involving walks, at-bats, hits, home runs, RBI, "
        "games, runs, stolen bases, strikeouts, average, OBP, SLG, or OPS."
    )
    inputs = {
        "team": {
            "type": "string",
            "description": "Team name, city, nickname, or abbreviation.",
        },
        "season": {
            "type": "integer",
            "description": "Four-digit MLB season.",
        },
        "sort_stat": {
            "type": "string",
            "description": "Statistic to rank, such as walks, at bats, hits, or home runs.",
        },
    }
    output_type = "string"

    def forward(self, team: str, season: int, sort_stat: str) -> str:
        aliases = {
            "walk": "baseOnBalls",
            "walks": "baseOnBalls",
            "bb": "baseOnBalls",
            "at bat": "atBats",
            "at bats": "atBats",
            "ab": "atBats",
            "hit": "hits",
            "hits": "hits",
            "home run": "homeRuns",
            "home runs": "homeRuns",
            "hr": "homeRuns",
            "rbi": "rbi",
            "games": "gamesPlayed",
            "runs": "runs",
            "stolen bases": "stolenBases",
            "strikeouts": "strikeOuts",
            "average": "avg",
            "avg": "avg",
            "obp": "obp",
            "slg": "slg",
            "ops": "ops",
        }
        requested = " ".join(str(sort_stat or "").lower().split())
        stat_field = aliases.get(requested)
        if not stat_field:
            return "Unsupported MLB statistic: " + requested

        try:
            season = int(season)
            teams_response = requests.get(
                "https://statsapi.mlb.com/api/v1/teams",
                params={"sportId": 1, "season": season},
                timeout=(WEBPAGE_CONNECT_TIMEOUT, WEBPAGE_READ_TIMEOUT),
            )
            teams_response.raise_for_status()
            query = re.sub(r"[^a-z0-9]", "", str(team).lower())
            scored_teams = []
            for item in teams_response.json().get("teams", []):
                names = [
                    item.get("name"),
                    item.get("teamName"),
                    item.get("clubName"),
                    item.get("locationName"),
                    item.get("abbreviation"),
                ]
                normalized = [
                    re.sub(r"[^a-z0-9]", "", str(value).lower())
                    for value in names
                    if value
                ]
                score = max(
                    (
                        3 if query == value else
                        2 if query and query in value else
                        1 if value and value in query else 0
                    )
                    for value in normalized
                )
                if score:
                    scored_teams.append((score, item))
            if not scored_teams:
                return f"No MLB team matched '{team}' in {season}."
            selected = max(scored_teams, key=lambda item: item[0])[1]

            stats_response = requests.get(
                "https://statsapi.mlb.com/api/v1/stats",
                params={
                    "stats": "season",
                    "group": "hitting",
                    "season": season,
                    "teamId": selected["id"],
                    "playerPool": "ALL",
                    "limit": 200,
                },
                timeout=(WEBPAGE_CONNECT_TIMEOUT, WEBPAGE_READ_TIMEOUT),
            )
            stats_response.raise_for_status()
            stats_groups = stats_response.json().get("stats", [])
            splits = stats_groups[0].get("splits", []) if stats_groups else []
            ranked = [
                split
                for split in splits
                if split.get("stat", {}).get(stat_field) not in (None, "")
            ]
            ranked.sort(
                key=lambda split: float(
                    str(split["stat"][stat_field]).replace(",", "")
                ),
                reverse=True,
            )
            ranked = ranked[:10]
            person_ids = ",".join(
                str(split.get("player", {}).get("id"))
                for split in ranked
                if split.get("player", {}).get("id")
            )
            names = {}
            if person_ids:
                people_response = requests.get(
                    "https://statsapi.mlb.com/api/v1/people",
                    params={"personIds": person_ids},
                    timeout=(WEBPAGE_CONNECT_TIMEOUT, WEBPAGE_READ_TIMEOUT),
                )
                people_response.raise_for_status()
                names = {
                    person["id"]: person.get("fullName", str(person["id"]))
                    for person in people_response.json().get("people", [])
                }

            rows = []
            for split in ranked:
                player_id = split.get("player", {}).get("id")
                stat = split.get("stat", {})
                rows.append(
                    f"{names.get(player_id, player_id)}: "
                    f"{stat_field}={stat.get(stat_field)}, "
                    f"atBats={stat.get('atBats')}, "
                    f"baseOnBalls={stat.get('baseOnBalls')}"
                )
            return (
                f"Official MLB season stats — {selected['name']} {season}, "
                f"ranked by {stat_field}:\n" + "\n".join(rows)
            )
        except Exception as exc:
            return f"MLB Stats API query failed: {compact_error(exc)}"


class TranscribeGaiaAudioTool(Tool):
    name = "transcribe_gaia_audio"
    description = (
        "Downloads and transcribes the official GAIA audio attachment. Use it "
        "first whenever the attachment is MP3, WAV, FLAC, M4A, OGG, or WEBM. "
        "It uses Hugging Face speech recognition, not the chat model."
    )
    inputs = {
        "task_id": {
            "type": "string",
            "description": "Exact GAIA task_id associated with the audio.",
        }
    }
    output_type = "string"

    def forward(self, task_id: str) -> str:
        task_id = str(task_id or "").strip()
        try:
            data, filename = download_gaia_attachment(task_id)
            suffix = Path(filename).suffix.lower()
            if suffix not in {
                ".mp3", ".wav", ".flac", ".m4a", ".ogg", ".webm", ".mp4"
            }:
                return (
                    f"Attachment {filename} is not an audio file. Use "
                    "inspect_gaia_attachment."
                )

            hf_token = os.getenv("HF_TOKEN")
            if not hf_token:
                return "Audio transcription failed: HF_TOKEN is not configured."
            try:
                from huggingface_hub import InferenceClient

                client = InferenceClient(api_key=hf_token, provider="auto")
                transcript_result = client.automatic_speech_recognition(
                    data,
                    model=os.getenv(
                        "GAIA_ASR_MODEL", "openai/whisper-large-v3"
                    ),
                )
                transcript = str(
                    getattr(transcript_result, "text", transcript_result)
                ).strip()
                if not transcript:
                    return "Hugging Face returned an empty audio transcript."
                return (
                    f"Audio transcript ({filename}):\n"
                    f"{transcript[:6_000]}"
                )
            except Exception as exc:
                return (
                    "Hugging Face speech-to-text failed: "
                    f"{compact_error(exc)}"
                )
        except Exception as exc:
            return (
                f"Could not transcribe audio for task {task_id}: "
                f"{compact_error(exc)}"
            )


class CalculatorTool(Tool):
    name = "calculator"
    description = (
        "Evaluates arithmetic locally without an LLM. Supports +, -, *, /, //, "
        "**, %, parentheses, pi, e, sqrt, log, exp, sin, cos, tan, abs, and round."
    )
    inputs = {
        "expression": {
            "type": "string",
            "description": "Arithmetic expression to evaluate.",
        }
    }
    output_type = "string"

    def forward(self, expression: str) -> str:
        binary_ops = {
            ast.Add: operator.add,
            ast.Sub: operator.sub,
            ast.Mult: operator.mul,
            ast.Div: operator.truediv,
            ast.FloorDiv: operator.floordiv,
            ast.Mod: operator.mod,
            ast.Pow: operator.pow,
        }
        unary_ops = {ast.UAdd: operator.pos, ast.USub: operator.neg}
        functions = {
            "abs": abs,
            "round": round,
            "sqrt": math.sqrt,
            "log": math.log,
            "exp": math.exp,
            "sin": math.sin,
            "cos": math.cos,
            "tan": math.tan,
        }
        constants = {"pi": math.pi, "e": math.e}

        def evaluate(node):
            if isinstance(node, ast.Expression):
                return evaluate(node.body)
            if isinstance(node, ast.Constant) and isinstance(
                node.value, (int, float)
            ):
                return node.value
            if isinstance(node, ast.BinOp) and type(node.op) in binary_ops:
                return binary_ops[type(node.op)](
                    evaluate(node.left), evaluate(node.right)
                )
            if isinstance(node, ast.UnaryOp) and type(node.op) in unary_ops:
                return unary_ops[type(node.op)](evaluate(node.operand))
            if isinstance(node, ast.Name) and node.id in constants:
                return constants[node.id]
            if (
                isinstance(node, ast.Call)
                and isinstance(node.func, ast.Name)
                and node.func.id in functions
                and not node.keywords
            ):
                return functions[node.func.id](
                    *(evaluate(argument) for argument in node.args)
                )
            raise ValueError("Unsupported expression.")

        try:
            parsed = ast.parse(str(expression), mode="eval")
            result = evaluate(parsed)
            return str(result)
        except Exception as exc:
            return f"Calculation failed: {compact_error(exc)}"


class QueryGaiaSpreadsheetTool(Tool):
    name = "query_gaia_spreadsheet"
    description = (
        "Analyzes an attached XLSX, XLSM, CSV, or TSV locally. Use operation "
        "'describe' first, then sum, mean, min, max, count, unique, or rows. "
        "This avoids sending the entire spreadsheet to the language model."
    )
    inputs = {
        "task_id": {
            "type": "string",
            "description": "Exact GAIA task_id for the spreadsheet.",
        },
        "operation": {
            "type": "string",
            "description": "describe, sum, mean, min, max, count, unique, or rows.",
        },
        "sheet": {
            "type": "string",
            "description": "Sheet name, or an empty string for the first sheet.",
        },
        "column": {
            "type": "string",
            "description": "Target column, or empty for describe/count rows.",
        },
        "filters": {
            "type": "string",
            "description": (
                "Optional exact filters as column=value;column=value. "
                "Use an empty string for no filter."
            ),
        },
    }
    output_type = "string"

    def forward(
        self,
        task_id: str,
        operation: str,
        sheet: str,
        column: str,
        filters: str,
    ) -> str:
        try:
            data, filename = download_gaia_attachment(str(task_id).strip())
            suffix = Path(filename).suffix.lower()
            if suffix in {".xlsx", ".xlsm"}:
                tables = pd.read_excel(BytesIO(data), sheet_name=None)
            elif suffix in {".csv", ".tsv"}:
                separator = "\t" if suffix == ".tsv" else ","
                tables = {
                    "data": pd.read_csv(
                        BytesIO(data),
                        sep=separator,
                        encoding_errors="replace",
                    )
                }
            else:
                return f"Attachment {filename} is not a supported spreadsheet."

            requested_sheet = str(sheet or "").strip()
            sheet_name = next(iter(tables))
            if requested_sheet:
                matching_sheet = next(
                    (
                        name
                        for name in tables
                        if str(name).lower() == requested_sheet.lower()
                    ),
                    None,
                )
                if matching_sheet is None:
                    return (
                        f"Unknown sheet {requested_sheet}. Available: "
                        + ", ".join(map(str, tables))
                    )
                sheet_name = matching_sheet

            frame = tables[sheet_name].copy()
            frame.columns = [str(value).strip() for value in frame.columns]
            requested_column = str(column or "").strip()

            for filter_expression in str(filters or "").split(";"):
                filter_expression = filter_expression.strip()
                if not filter_expression:
                    continue
                if "=" not in filter_expression:
                    return f"Invalid filter: {filter_expression}"
                filter_column, filter_value = (
                    part.strip() for part in filter_expression.split("=", 1)
                )
                actual_filter_column = next(
                    (
                        name
                        for name in frame.columns
                        if name.lower() == filter_column.lower()
                    ),
                    None,
                )
                if actual_filter_column is None:
                    return (
                        f"Unknown filter column {filter_column}. Columns: "
                        + ", ".join(frame.columns)
                    )
                numeric_value = pd.to_numeric(
                    pd.Series([filter_value]), errors="coerce"
                ).iloc[0]
                numeric_column = pd.to_numeric(
                    frame[actual_filter_column], errors="coerce"
                )
                if pd.notna(numeric_value) and numeric_column.notna().any():
                    frame = frame[numeric_column == numeric_value]
                else:
                    frame = frame[
                        frame[actual_filter_column]
                        .astype(str)
                        .str.strip()
                        .str.casefold()
                        == filter_value.casefold()
                    ]

            operation = str(operation or "describe").strip().lower()
            if operation == "describe":
                preview = frame.head(5).to_csv(index=False)
                return (
                    f"Workbook: {filename}\n"
                    f"Sheets: {', '.join(map(str, tables))}\n"
                    f"Selected sheet: {sheet_name}\n"
                    f"Rows: {len(frame)}; Columns: {len(frame.columns)}\n"
                    f"Column names: {', '.join(frame.columns)}\n"
                    f"First rows:\n{preview[:2_500]}"
                )

            if requested_column:
                actual_column = next(
                    (
                        name
                        for name in frame.columns
                        if name.lower() == requested_column.lower()
                    ),
                    None,
                )
                if actual_column is None:
                    return (
                        f"Unknown target column {requested_column}. Columns: "
                        + ", ".join(frame.columns)
                    )
            else:
                actual_column = ""

            if operation == "count":
                result = (
                    int(frame[actual_column].notna().sum())
                    if actual_column
                    else len(frame)
                )
            elif operation in {"sum", "mean", "min", "max"}:
                if not actual_column:
                    return f"Operation {operation} requires a target column."
                series = pd.to_numeric(frame[actual_column], errors="coerce").dropna()
                if series.empty:
                    return f"Column {actual_column} has no numeric values."
                result = getattr(series, operation)()
            elif operation == "unique":
                if not actual_column:
                    return "Operation unique requires a target column."
                values = frame[actual_column].dropna().astype(str).unique().tolist()
                return ", ".join(values[:100])
            elif operation == "rows":
                columns = [actual_column] if actual_column else list(frame.columns)
                return frame[columns].head(30).to_csv(index=False)[:4_000]
            else:
                return (
                    "Unknown operation. Use describe, sum, mean, min, max, "
                    "count, unique, or rows."
                )

            if hasattr(result, "item"):
                result = result.item()
            return (
                f"Operation: {operation}; sheet: {sheet_name}; "
                f"rows matched: {len(frame)}; result: {result}"
            )
        except Exception as exc:
            return f"Spreadsheet analysis failed: {compact_error(exc)}"


class InspectGaiaAttachmentTool(Tool):
    name = "inspect_gaia_attachment"
    description = (
        "Downloads and reads the official attachment associated with a GAIA "
        "task. Use it for PDF, DOCX, text, code, or ZIP files. Do not use it "
        "for audio, images, or spreadsheets; those have specialized tools."
    )
    inputs = {
        "task_id": {
            "type": "string",
            "description": "The exact GAIA task_id supplied in the user task.",
        },
        "query": {
            "type": "string",
            "description": "Short terms for the exact fact to extract.",
        },
    }
    output_type = "string"

    def forward(self, task_id: str, query: str) -> str:
        task_id = str(task_id).strip()
        if not task_id:
            return "No task_id was supplied."

        try:
            data, filename = download_gaia_attachment(task_id)
            suffix = Path(filename).suffix.lower()
            if suffix in {
                ".mp3", ".wav", ".flac", ".m4a", ".ogg", ".webm", ".mp4"
            }:
                return (
                    f"Attachment {filename} is audio. Call "
                    "transcribe_gaia_audio with this task_id."
                )
            if suffix in {".png", ".jpg", ".jpeg", ".webp", ".gif"}:
                return (
                    f"Attachment {filename} is an image. Call "
                    "analyze_gaia_image with this task_id and question."
                )
            if suffix in {".xlsx", ".xlsm", ".csv", ".tsv"}:
                return (
                    f"Attachment {filename} is tabular. Call "
                    "query_gaia_spreadsheet with operation='describe' first."
                )
            extracted = extract_attachment_text(
                data, filename, query=str(query or "")
            )
            return f"Attachment filename: {filename}\n\n{extracted}"
        except Exception as exc:
            return (
                f"Could not inspect attachment for task {task_id}: "
                f"{compact_error(exc)}"
            )


class YouTubeTranscriptTool(Tool):
    name = "youtube_transcript"
    description = (
        "Retrieves the spoken transcript or subtitles of a YouTube video. "
        "Use it for questions asking what a person says in a linked video. "
        "It cannot determine purely visual events."
    )
    inputs = {
        "url": {
            "type": "string",
            "description": "Full YouTube URL or the 11-character video ID.",
        },
        "query": {
            "type": "string",
            "description": "Short terms describing the spoken fact or quote.",
        },
    }
    output_type = "string"

    def forward(self, url: str, query: str) -> str:
        from youtube_transcript_api import YouTubeTranscriptApi

        value = str(url or "").strip()
        match = re.search(
            r"(?:v=|youtu\.be/|shorts/)([A-Za-z0-9_-]{11})", value
        )
        video_id = match.group(1) if match else value
        if not re.fullmatch(r"[A-Za-z0-9_-]{11}", video_id):
            return "Could not identify a valid YouTube video ID."

        try:
            api = YouTubeTranscriptApi()
            transcript = api.fetch(video_id)
            lines = []
            for snippet in transcript:
                text = getattr(snippet, "text", None)
                if text is None and isinstance(snippet, dict):
                    text = snippet.get("text")
                if text:
                    lines.append(str(text))
            result = " ".join(lines).strip()
            return (
                focus_text(result, str(query or ""), 5_000)
                if result
                else "The video has no available transcript."
            )
        except Exception as exc:
            return (
                "Could not retrieve YouTube transcript: "
                f"{compact_error(exc)}"
            )


class AnalyzeGaiaImageTool(Tool):
    name = "analyze_gaia_image"
    description = (
        "Downloads the official image for a GAIA task and analyzes it with a "
        "vision model. Use this for questions whose answer depends on image "
        "pixels, diagrams, chess positions, or visual details."
    )
    inputs = {
        "task_id": {
            "type": "string",
            "description": "The exact GAIA task_id associated with the image.",
        },
        "question": {
            "type": "string",
            "description": "The complete question the image must answer.",
        },
    }
    output_type = "string"

    def forward(self, task_id: str, question: str) -> str:
        token = os.getenv("HF_TOKEN")
        if not token:
            return "Image analysis failed: HF_TOKEN is not configured."

        try:
            data, filename = download_gaia_attachment(str(task_id).strip())
            suffix = Path(filename).suffix.lower()
            mime = {
                ".png": "image/png",
                ".jpg": "image/jpeg",
                ".jpeg": "image/jpeg",
                ".webp": "image/webp",
                ".gif": "image/gif",
            }.get(suffix, "image/png")
            encoded = base64.b64encode(data).decode("ascii")

            vision_model = os.getenv(
                "GAIA_VISION_MODEL",
                "Qwen/Qwen3-VL-235B-A22B-Instruct:cheapest",
            )
            response = requests.post(
                "https://router.huggingface.co/v1/chat/completions",
                headers={
                    "Authorization": f"Bearer {token}",
                    "Content-Type": "application/json",
                },
                json={
                    "model": vision_model,
                    "messages": [
                        {
                            "role": "user",
                            "content": [
                                {
                                    "type": "text",
                                    "text": (
                                        "Analyze the supplied image carefully and "
                                        "answer this task. Return the likely exact "
                                        "answer plus at most one short evidence "
                                        "sentence:\n"
                                        f"{question}"
                                    ),
                                },
                                {
                                    "type": "image_url",
                                    "image_url": {
                                        "url": f"data:{mime};base64,{encoded}"
                                    },
                                },
                            ],
                        }
                    ],
                    "temperature": 0,
                    "max_tokens": 300,
                },
                timeout=HTTP_TIMEOUT,
            )
            response.raise_for_status()
            payload = response.json()
            return str(payload["choices"][0]["message"]["content"]).strip()
        except Exception as exc:
            return (
                "Could not analyze the GAIA image: "
                f"{compact_error(exc)}"
            )


class BasicAgent:
    def __init__(self):
        print("Inicializando o agente GAIA...")

        hf_token = os.getenv("HF_TOKEN")
        gemini_api_key = os.getenv("GEMINI_API_KEY")
        configured_model = str(
            os.getenv("GAIA_GEMINI_MODEL") or DEFAULT_GEMINI_MODEL
        ).strip()
        model_name = configured_model.removeprefix("gemini/")
        if not model_name.lower().startswith("gemini-"):
            print(
                "GAIA_GEMINI_MODEL não apontava para um modelo Gemini "
                "e foi ignorado. "
                f"Usando {DEFAULT_GEMINI_MODEL}."
            )
            model_name = DEFAULT_GEMINI_MODEL
        model_id = f"gemini/{model_name}"
        if not gemini_api_key:
            raise RuntimeError(
                "O secret GEMINI_API_KEY não está configurado. "
                "Adicione a chave em Settings > Variables and secrets > Secrets."
            )

        self.model = LiteLLMModel(
            model_id=model_id,
            api_key=gemini_api_key,
            max_tokens=1_200,
            reasoning_effort="low",
            requests_per_minute=8,
        )
        self.hf_token = hf_token
        self.model_id = model_id
        self.run_tracker = RunEvidenceTracker()
        configured_max_steps = str(
            os.getenv("GAIA_MAX_STEPS") or "unlimited"
        ).strip()
        if configured_max_steps.lower() in {
            "",
            "none",
            "unlimited",
            "infinite",
            "infinity",
        }:
            self.max_steps = PRACTICALLY_UNLIMITED_STEPS
            steps_label = "sem limite prático"
        else:
            try:
                self.max_steps = max(1, int(configured_max_steps))
            except ValueError:
                self.max_steps = PRACTICALLY_UNLIMITED_STEPS
                steps_label = "sem limite prático"
            else:
                steps_label = str(self.max_steps)
        print(f"Modelo principal selecionado: {model_id}")
        print(f"Limite de passos do agente: {steps_label}")

        web_search_tool = ConciseWebSearchTool()
        visit_page_tool = OpenWebPageTool()
        wikipedia_tool = WikipediaSearchTool(
            user_agent="GAIA-Course-Agent/1.0 (educational project)",
            language="en",
        )
        wikipedia_tool.description = (
            "Searches current English Wikipedia content directly. Use it for "
            "encyclopedic facts when no historical version is requested. For "
            "a version-specific year, use wikipedia_revision instead."
        )

        agent_tools = [
            web_search_tool,
            visit_page_tool,
            QueryWebTableTool(),
            MlbStatsTool(),
            ReadDocumentUrlTool(),
            WikipediaRevisionTool(),
            WikipediaFeaturedArticlesTool(),
            wikipedia_tool,
            InspectGaiaAttachmentTool(),
            QueryGaiaSpreadsheetTool(),
            TranscribeGaiaAudioTool(),
            YouTubeTranscriptTool(),
            AnalyzeGaiaImageTool(),
            CalculatorTool(),
        ]
        agent_tools = [
            instrument_tool(tool, self.run_tracker)
            for tool in agent_tools
        ]
        self.tools_by_name = {tool.name: tool for tool in agent_tools}

        # Gemini returns native function calls. ToolCallingAgent handles that
        # structured format without parsing generated Python code.
        self.agent = ToolCallingAgent(
            tools=agent_tools,
            model=self.model,
            max_steps=self.max_steps,
            max_tool_threads=1,
            planning_interval=None,
            description=(
                "Agent designed to solve GAIA benchmark questions with "
                "exact-match answers."
            ),
        )

        exact_match_prompt = """
You are an expert AI assistant solving tasks from the GAIA benchmark.

TOOL ROUTING POLICY:
1. web_search discovers URLs and snippets. It does not read full pages.
2. visit_webpage reads focused passages from HTML; always provide short query
   terms. query_web_table computes min/max from an HTML table locally; use it
   instead of reading a long statistics table. read_document_url reads linked
   PDFs or documents by query. mlb_stats queries official baseball season data.
3. transcribe_gaia_audio is the only tool for attached audio. Call it first
   for MP3, WAV, FLAC, M4A, OGG, WEBM, or MP4.
4. inspect_gaia_attachment handles attached PDF, DOCX, text, code, or ZIP.
   Pass short target terms. For XLSX/CSV use
   query_gaia_spreadsheet instead: describe columns, then run one calculation.
   Never send a whole table to the model.
5. analyze_gaia_image handles attached images. youtube_transcript handles
   spoken YouTube content. calculator evaluates arithmetic locally.
6. wikipedia_search is for current Wikipedia or encyclopedic facts.
   wikipedia_revision is mandatory when the task specifies a Wikipedia
   version/year; it preserves historical table rows omitted by summaries.
   wikipedia_featured_articles is mandatory for Featured Article promotion
   questions with a month/year; it returns the official article/nominator rows.
   Verify other historical, nomination, or archive details with an exact page.

Research carefully, prefer primary or official sources, and cross-check
uncertain facts. A search snippet alone is insufficient when the source page
can be opened. Never invent a tool, use subprocess, or use shell commands.
When CONTROLLER-PRECOLLECTED EVIDENCE is present, use it first and do not
repeat its exact tool call. If it directly answers the task, immediately call
final_answer instead of researching again.
Do not repeat nearly identical searches. Stop as soon as primary evidence
answers the exact question. Never pass an entire task as a webpage query; use
only names, identifiers, and the target fact. Continue working until the
requested value is supported, then call final_answer immediately.

FINAL RESPONSE POLICY:
Call the final_answer tool with only the requested value. Never write the
answer as plain text instead of calling final_answer. Never
include reasoning, explanations, labels, Markdown, citations, or the words
"FINAL ANSWER" inside the submitted value.
- Quantity/count: return only the number, unless units or currency are requested.
- Person: return only the requested name component.
- City/country/code: return only that value.
- List: return only items with the requested separator and ordering.
- Chess move: return only algebraic notation.
- Quote: return only the requested spoken words.
"""
        self.agent.prompt_templates["system_prompt"] = (
            exact_match_prompt.strip()
            + "\n\n"
            + self.agent.prompt_templates["system_prompt"]
        )
        self.run_tracker.interrupt_callback = self.agent.interrupt

    def interrupt(self):
        """Interrompe com segurança a execução atual do smolagents."""
        self.agent.interrupt()

    def _precollect_deterministic_evidence(
        self,
        question: str,
        task_id: str | None,
        attachment_name: str = "",
    ) -> str:
        """Runs mandatory/specialized tools by rule before an LLM can choose."""
        suffix = Path(attachment_name).suffix.lower()
        query = concise_query_from_question(question)
        route = ""

        if task_id and suffix in {
            ".mp3", ".wav", ".flac", ".m4a", ".ogg", ".webm", ".mp4"
        }:
            self.tools_by_name["transcribe_gaia_audio"].forward(task_id=task_id)
            route = "audio attachment -> transcribe_gaia_audio"
        elif task_id and suffix in {
            ".png", ".jpg", ".jpeg", ".webp", ".gif"
        }:
            self.tools_by_name["analyze_gaia_image"].forward(
                task_id=task_id,
                question=question,
            )
            route = "image attachment -> analyze_gaia_image"
        elif task_id and suffix in {".xlsx", ".xlsm", ".csv", ".tsv"}:
            self.tools_by_name["query_gaia_spreadsheet"].forward(
                task_id=task_id,
                operation="describe",
                sheet="",
                column="",
                filters="",
            )
            route = "spreadsheet attachment -> query_gaia_spreadsheet(describe)"
        elif task_id and attachment_name:
            self.tools_by_name["inspect_gaia_attachment"].forward(
                task_id=task_id,
                query=query,
            )
            route = "document attachment -> inspect_gaia_attachment"

        youtube_match = re.search(
            r"https?://(?:www\.)?(?:youtube\.com/watch\?v=|youtu\.be/)"
            r"[A-Za-z0-9_-]{11}",
            question,
            flags=re.I,
        )
        spoken_cues = (
            "say", "said", "says", "speak", "spoken", "quote", "transcript",
            "according to the video", "what does", "what did",
        )
        if (
            not route
            and youtube_match
            and any(cue in question.lower() for cue in spoken_cues)
        ):
            self.tools_by_name["youtube_transcript"].forward(
                url=youtube_match.group(0),
                query=query,
            )
            route = "spoken YouTube question -> youtube_transcript"

        team_aliases = {
            "yankee": "Yankees", "red sox": "Red Sox",
            "oriole": "Orioles", "ray": "Rays", "blue jay": "Blue Jays",
            "white sox": "White Sox", "guardian": "Guardians",
            "cleveland indian": "Indians", "tiger": "Tigers",
            "royal": "Royals", "twin": "Twins", "astro": "Astros",
            "angel": "Angels", "athletic": "Athletics", "mariner": "Mariners",
            "ranger": "Rangers", "brave": "Braves", "marlin": "Marlins",
            "met": "Mets", "phillie": "Phillies", "national": "Nationals",
            "cub": "Cubs", "red": "Reds", "brewer": "Brewers",
            "pirate": "Pirates", "cardinal": "Cardinals",
            "diamondback": "Diamondbacks", "rockie": "Rockies",
            "dodger": "Dodgers", "padre": "Padres", "giant": "Giants",
        }
        lower_question = question.lower()
        years = re.findall(r"\b(?:18|19|20)\d{2}\b", question)

        month_match = re.search(
            r"\b(January|February|March|April|May|June|July|August|"
            r"September|October|November|December)\b",
            question,
            flags=re.I,
        )
        featured_subject_match = re.search(
            r"\babout\s+(?:an?\s+|the\s+)?(.+?)\s+"
            r"(?:that|which)\s+was\s+promoted\b",
            question,
            flags=re.I,
        )
        if (
            not route
            and "featured article" in lower_question
            and "promoted" in lower_question
            and years
            and month_match
        ):
            subject = (
                featured_subject_match.group(1).strip(" ?.,")
                if featured_subject_match
                else ""
            )
            self.tools_by_name["wikipedia_featured_articles"].forward(
                year=int(years[0]),
                month=month_match.group(1),
                subject=subject,
            )
            route = (
                "Wikipedia Featured Article promotion -> official monthly "
                "promotion table"
            )

        matched_team = next(
            (
                canonical
                for alias, canonical in team_aliases.items()
                if re.search(rf"\b{re.escape(alias)}s?\b", lower_question)
            ),
            None,
        )
        stat_cues = [
            ("most walks", "walks"), ("least walks", "walks"),
            ("walks", "walks"), ("at bats", "at bats"),
            ("home runs", "home runs"), ("stolen bases", "stolen bases"),
            ("strikeouts", "strikeouts"), ("hits", "hits"),
            ("runs batted", "rbi"), ("rbi", "rbi"), ("ops", "ops"),
            ("obp", "obp"), ("slugging", "slg"), ("average", "average"),
        ]
        matched_stat = next(
            (stat for cue, stat in stat_cues if cue in lower_question),
            None,
        )
        if not route and matched_team and years and matched_stat:
            self.tools_by_name["mlb_stats"].forward(
                team=matched_team,
                season=int(years[0]),
                sort_stat=matched_stat,
            )
            route = "MLB statistics question -> official MLB Stats API"

        if route:
            print(f"Rota determinística: {route}")
        return route

    @staticmethod
    def _invalid_candidate(candidate: str) -> bool:
        value = str(candidate or "").strip()
        lowered = value.lower()
        if not value or lowered in {"none", "null", "n/a"}:
            return True
        if len(value) > 1_000:
            return True
        tool_syntax = (
            bool(re.search(
                r"""["']type["']\s*:\s*["']function["']""",
                value,
                flags=re.I,
            ))
            and bool(re.search(
                r"""["']arguments["']\s*:""",
                value,
                flags=re.I,
            ))
        )
        unfinished_cues = (
            "call: ", "calling tool", "let's search", "lets search",
            "let's do a search", "web_search(", "wikipedia_search(",
            "visit_webpage(", "query_web_table(", "i need to search",
        )
        return tool_syntax or any(cue in lowered for cue in unfinished_cues)

    def _gemini_answer_from_evidence(
        self,
        question: str,
        evidence: str,
        task_id: str | None = None,
    ) -> str:
        """Produces one answer from existing evidence without running tools."""
        gemini_api_key = os.getenv("GEMINI_API_KEY")
        if not gemini_api_key:
            raise RuntimeError(
                "GEMINI_API_KEY não está configurada."
            )
        model = os.getenv(
            "GAIA_GEMINI_MODEL", DEFAULT_GEMINI_MODEL
        )
        model = str(model).removeprefix("gemini/")
        if not model.lower().startswith("gemini-"):
            model = DEFAULT_GEMINI_MODEL
        prompt = f"""
Answer this GAIA task using ONLY the collected evidence below. Do not call or
suggest tools and do not perform another search. If the evidence is sufficient,
return only the exact requested value in final_answer. If it is insufficient,
set sufficient_evidence to false and leave final_answer empty.

Task ID: {task_id or "test"}
Question: {question}

COLLECTED EVIDENCE:
{evidence}
""".strip()
        response = requests.post(
            (
                "https://generativelanguage.googleapis.com/v1beta/models/"
                f"{model}:generateContent"
            ),
            headers={
                "x-goog-api-key": gemini_api_key,
                "Content-Type": "application/json",
            },
            json={
                "contents": [
                    {"role": "user", "parts": [{"text": prompt}]}
                ],
                "generationConfig": {
                    "maxOutputTokens": 512,
                    "thinkingConfig": {"thinkingLevel": "minimal"},
                    "responseMimeType": "application/json",
                    "responseJsonSchema": {
                        "type": "object",
                        "properties": {
                            "final_answer": {"type": "string"},
                            "sufficient_evidence": {"type": "boolean"},
                        },
                        "required": [
                            "final_answer",
                            "sufficient_evidence",
                        ],
                        "additionalProperties": False,
                    },
                },
            },
            timeout=(WEBPAGE_CONNECT_TIMEOUT, 45),
        )
        response.raise_for_status()
        payload = response.json()
        parts = payload["candidates"][0]["content"]["parts"]
        content = "".join(
            str(part.get("text") or "") for part in parts
        ).strip()
        data = json.loads(content)
        if not data.get("sufficient_evidence"):
            raise RuntimeError(
                "O Gemini informou que as evidências coletadas ainda são "
                "insuficientes. A resposta não foi salva."
            )
        answer = self.enforce_direct_answer(
            question, str(data.get("final_answer") or "")
        )
        if self._invalid_candidate(answer):
            raise RuntimeError(
                "O Gemini não produziu uma resposta final válida."
            )
        print(f"Gemini respondeu usando as evidências existentes: {answer}")
        return answer

    def _run_gemini(self, task_context: str):
        try:
            return self.agent.run(task_context, reset=True)
        except Exception as exc:
            if self.run_tracker.stagnated:
                raise ResearchStagnationError(
                    self.run_tracker.stagnation_reason
                ) from exc
            error_text = str(exc).lower()
            retryable_provider_error = any(
                marker in error_text
                for marker in (
                    "payment required",
                    "insufficient credit",
                    "billing",
                    "quota",
                    "rate limit",
                    "429",
                    "authenticationerror",
                    "unauthorized",
                    "401",
                    "service unavailable",
                    "timeout",
                )
            )
            if retryable_provider_error:
                raise RuntimeError(
                    "O Gemini recusou a chamada. Verifique GEMINI_API_KEY, "
                    "a cota e os limites do projeto no Google AI Studio. "
                    f"Detalhe: {compact_error(exc)}"
                ) from exc
            raise

    def __call__(self, question: str, task_id: str | None = None) -> str:
        question = (question or "").strip()
        if not question:
            raise ValueError("Digite uma pergunta para testar o agente.")
        SEARCH_CACHE.clear()
        self.run_tracker.reset(question)
        attachment_name = ""

        if task_id:
            attachment_name = get_task_file_name(task_id)
            attachment_suffix = Path(attachment_name).suffix.lower()
            if attachment_suffix in {
                ".mp3", ".wav", ".flac", ".m4a", ".ogg", ".webm", ".mp4"
            }:
                attachment_context = (
                    f"Official attachment: {attachment_name}. Call "
                    "transcribe_gaia_audio first."
                )
            elif attachment_suffix in {
                ".png", ".jpg", ".jpeg", ".webp", ".gif"
            }:
                attachment_context = (
                    f"Official attachment: {attachment_name}. Call "
                    "analyze_gaia_image with this task_id and question."
                )
            elif attachment_suffix in {".xlsx", ".xlsm", ".csv", ".tsv"}:
                attachment_context = (
                    f"Official attachment: {attachment_name}. Call "
                    "query_gaia_spreadsheet with operation='describe' first."
                )
            elif attachment_name:
                attachment_context = (
                    f"Official attachment: {attachment_name}. Call "
                    "inspect_gaia_attachment with this task_id and short target terms."
                )
            else:
                attachment_context = (
                    "Official attachment: NONE. Do not call any GAIA "
                    "attachment tool."
                )
            task_context = (
                f"GAIA task_id: {task_id}\n"
                f"{attachment_context}\n\nQuestion: {question}"
            )
        else:
            task_context = question
        route = self._precollect_deterministic_evidence(
            question=question,
            task_id=task_id,
            attachment_name=attachment_name,
        )
        precollected = self.run_tracker.current()
        if precollected:
            task_context += (
                "\n\nCONTROLLER-PRECOLLECTED EVIDENCE:\n"
                f"{precollected}\n\n"
                f"Controller route: {route}. Do not repeat the same tool call; "
                "a new spreadsheet calculation is allowed after describe. "
                "If this evidence answers the question, call final_answer now."
            )
        try:
            result = self._run_gemini(task_context)
        except ResearchStagnationError as exc:
            evidence = self.run_tracker.current()
            if not evidence:
                raise RuntimeError(
                    "A pesquisa estagnou sem produzir evidência utilizável."
                ) from exc
            print(
                "Pesquisa interrompida por estagnação; consolidando as "
                f"evidências. Motivo: {exc}"
            )
            result = self._gemini_answer_from_evidence(
                question=question,
                evidence=evidence,
                task_id=task_id,
            )
        except Exception as exc:
            error_text = str(exc)
            if (
                "AuthenticationError" in error_text
                or "401" in error_text
                or "Unauthorized" in error_text
            ):
                raise RuntimeError(
                    f"Falha de autenticação no modelo {self.model_id}. "
                    "Verifique o secret GEMINI_API_KEY."
                ) from exc
            raise
        candidate = self.enforce_direct_answer(question, str(result))
        invalid_candidate = (
            not candidate
            or candidate.strip().lower() in {"none", "null"}
            or (
                bool(re.search(
                    r"""["']type["']\s*:\s*["']function["']""",
                    candidate,
                    flags=re.I,
                ))
                and bool(re.search(
                    r"""["']arguments["']\s*:""",
                    candidate,
                    flags=re.I,
                ))
            )
        )
        invalid_candidate = (
            invalid_candidate or self._invalid_candidate(candidate)
        )
        if invalid_candidate and self.run_tracker.current():
            print(
                "Gemini terminou sem resposta final; tentando consolidar "
                "somente as evidências já coletadas."
            )
            evidence = self.run_tracker.current()
            try:
                candidate = self._gemini_answer_from_evidence(
                    question=question,
                    evidence=evidence,
                    task_id=task_id,
                )
            except Exception as exc:
                print(
                    "A consolidação Gemini falhou: "
                    f"{compact_error(exc)}"
                )
            invalid_candidate = (
                not candidate
                or candidate.strip().lower() in {"none", "null"}
                or (
                    bool(re.search(
                        r"""["']type["']\s*:\s*["']function["']""",
                        candidate,
                        flags=re.I,
                    ))
                    and bool(re.search(
                        r"""["']arguments["']\s*:""",
                        candidate,
                        flags=re.I,
                    ))
                )
            )
            invalid_candidate = (
                invalid_candidate or self._invalid_candidate(candidate)
            )
        if invalid_candidate:
            raise RuntimeError(
                "O agente não produziu uma resposta final válida. "
                "A resposta não foi salva; execute novamente esta questão."
            )
        return self.review_answer_with_gemini(
            question=question,
            candidate=candidate,
            task_id=task_id,
            evidence=self.run_tracker.current(),
        )

    @staticmethod
    def deterministic_answer_cleanup(answer: str) -> str:
        """Remove embalagens comuns sem alterar o conteúdo da resposta."""
        text = str(answer or "").strip()
        text = re.sub(r"</?code>", "", text, flags=re.I).strip()
        text = re.sub(r"^```(?:text|markdown)?\s*", "", text, flags=re.I)
        text = re.sub(r"\s*```$", "", text)
        text = re.sub(r"^\s*#{1,6}\s*", "", text)
        text = re.sub(r"^\s*[-*•]\s+", "", text)
        text = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", text)

        final_call = re.search(
            r"final_answer\s*\(\s*([\"']?)(.*?)\1\s*\)\s*$",
            text,
            flags=re.I | re.S,
        )
        if final_call:
            text = final_call.group(2).strip()

        marker_pattern = re.compile(
            r"(?:final\s+answer|answer|resposta\s+final|resposta)\s*:\s*",
            flags=re.I,
        )
        marker_matches = list(marker_pattern.finditer(text))
        if marker_matches:
            text = text[marker_matches[-1].end() :].strip()

        prefix_patterns = [
            r"^the\s+(?:final\s+)?answer\s+is\s+",
            r"^my\s+(?:final\s+)?answer\s+is\s+",
            r"^a\s+resposta(?:\s+final)?\s+[ée]\s+",
        ]
        for pattern in prefix_patterns:
            text = re.sub(pattern, "", text, flags=re.I).strip()

        text = text.replace("**", "").replace("__", "").strip()
        if (
            len(text) >= 2
            and text[0] == text[-1]
            and text[0] in {'"', "'", "`"}
        ):
            text = text[1:-1].strip()

        return text.replace("FINAL ANSWER", "").strip()

    @classmethod
    def enforce_direct_answer(cls, question: str, answer: str) -> str:
        """Impõe o formato exact-match sem pedir nova interpretação a uma LLM."""
        original_text = str(answer or "")
        bold_values = [
            value.strip()
            for value in re.findall(r"\*\*(.+?)\*\*", original_text, flags=re.S)
            if value.strip()
        ]
        text = cls.deterministic_answer_cleanup(answer)
        question_lower = str(question or "").lower()

        lines = [line.strip() for line in text.splitlines() if line.strip()]
        if len(lines) > 1:
            # Para listas, privilegia a linha que realmente contém os itens.
            if "comma" in question_lower or "vírgula" in question_lower:
                comma_lines = [line for line in lines if "," in line]
                if comma_lines:
                    text = max(comma_lines, key=lambda value: value.count(","))
                else:
                    text = lines[-1]
            else:
                text = lines[-1]

        text = cls.deterministic_answer_cleanup(text)

        if "page numbers" in question_lower or "page number" in question_lower:
            page_groups = re.findall(
                r"\b(?:pages?|pp\.?)\s*(?:are|is|:|-)?\s*"
                r"(\d+(?:(?:\s*,\s*(?:and\s+)?|\s+and\s+|-)\d+)*)",
                text,
                flags=re.I,
            )
            if page_groups:
                pages = [int(value) for value in re.findall(r"\d+", page_groups[-1])]
                if pages:
                    return ", ".join(str(value) for value in sorted(set(pages)))

        quantity_question = (
            "how many" in question_lower
            or "numeric output" in question_lower
            or "quantos" in question_lower
            or "quantas" in question_lower
        )
        if quantity_question:
            numbers = re.findall(
                r"(?<![\w.])-?\d+(?:,\d{3})*(?:\.\d+)?", text
            )
            if numbers:
                return numbers[-1].replace(",", "")

        requests_usd = (
            "in usd" in question_lower
            or "usd with" in question_lower
            or "dollars" in question_lower
        )
        if requests_usd:
            amounts = re.findall(
                r"\$?\s*(-?\d+(?:,\d{3})*(?:\.\d+)?)", text
            )
            if amounts:
                raw_amount = amounts[-1].replace(",", "")
                try:
                    return f"${float(raw_amount):,.2f}"
                except ValueError:
                    pass

        if "award number" in question_lower or "grant number" in question_lower:
            identifiers = re.findall(r"\b[A-Z0-9][A-Z0-9-]{5,}\b", text.upper())
            identifiers = [
                value
                for value in identifiers
                if re.search(r"[A-Z]", value) and re.search(r"\d", value)
            ]
            if identifiers:
                return identifiers[-1].strip(" .,:;\"'")

        list_question = (
            "comma" in question_lower
            or "list" in question_lower
            or "separated" in question_lower
            or "delimited" in question_lower
        )
        if (
            list_question
            and "last name" in question_lower
            and "before" in question_lower
            and "after" in question_lower
        ):
            emphasized_names = [
                cls.deterministic_answer_cleanup(value).strip(" .,:;\"'")
                for value in bold_values
            ]
            emphasized_names = [
                value
                for value in emphasized_names
                if value and len(value) <= 60 and " " not in value
            ]
            if len(emphasized_names) >= 2:
                return ", ".join(emphasized_names[-2:])

            name_pair = re.search(
                r"\b(?:are|were)\s+([A-Z][A-Za-z'’-]+)\s*"
                r"(?:,|and)\s*([A-Z][A-Za-z'’-]+)",
                text,
            )
            if name_pair:
                return f"{name_pair.group(1)}, {name_pair.group(2)}"

        person_question = not list_question and (
            question_lower.startswith("who ")
            or " who " in f" {question_lower} "
            or "first name" in question_lower
            or "surname" in question_lower
            or "username" in question_lower
        )
        if person_question:
            # Explanatory answers often repeat the requested person in the
            # final bold fragment. Prefer it before trying sentence patterns.
            if bold_values:
                emphasized = cls.deterministic_answer_cleanup(bold_values[-1])
                if (
                    emphasized
                    and len(emphasized) <= 100
                    and not re.search(r"[.!?]\s+\w", emphasized)
                ):
                    return emphasized.strip(" .,:;\"'")

            person_patterns = [
                r"\b(?:nominated|written|directed|created|founded|authored|performed)"
                r"\s+by\s+([A-Z][\w'’-]*(?:\s+[A-Z][\w'’-]*){0,3})",
                r"\b(?:username|first\s+name|surname|name)\s+(?:is|was)\s+"
                r"([A-Z][\w'’-]*(?:\s+[A-Z][\w'’-]*){0,3})",
            ]
            matches = []
            for pattern in person_patterns:
                matches.extend(re.findall(pattern, text))
            if matches:
                return matches[-1].strip(" .,:;\"'")

        # Remove frases introdutórias que ainda possam aparecer em uma linha.
        text = re.sub(
            r"^(?:therefore,\s*|thus,\s*|so,\s*)?"
            r"(?:the\s+)?(?:correct\s+|final\s+)?answer\s+is\s+",
            "",
            text,
            flags=re.I,
        ).strip()
        text = re.sub(
            r"^(?:the\s+requested\s+)?"
            r"(?:first\s+name|surname|city|country|ioc\s+code)\s+is\s+",
            "",
            text,
            flags=re.I,
        ).strip()

        # Se ainda restar uma explicação seguida de dois-pontos, conserva o valor.
        if ":" in text:
            prefix, value = text.rsplit(":", 1)
            if len(value.strip()) <= 250 and any(
                cue in prefix.lower()
                for cue in ("answer", "resposta", "result", "resultado")
            ):
                text = value.strip()

        return cls.deterministic_answer_cleanup(text)

    def format_exact_answer(self, question: str, raw_answer: str) -> str:
        """Limpa o resultado mecanicamente, sem pedir a outro modelo para alterá-lo."""
        del question
        cleaned = self.deterministic_answer_cleanup(raw_answer)
        lines = [line.strip() for line in cleaned.splitlines() if line.strip()]

        # Quando o agente ainda inclui uma explicação e deixa uma resposta curta
        # isolada na última linha, conserva somente essa última linha.
        if len(lines) > 1:
            last_line = lines[-1]
            reasoning_cues = (
                "because",
                "therefore",
                "research",
                "source",
                "conclude",
                "analysis",
                "porque",
                "portanto",
                "pesquisa",
                "conclu",
            )
            preceding = " ".join(lines[:-1]).lower()
            if len(last_line) <= 250 and any(
                cue in preceding for cue in reasoning_cues
            ):
                cleaned = last_line

        return self.deterministic_answer_cleanup(cleaned)

    def review_answer_with_gemini(
        self,
        question: str,
        candidate: str,
        task_id: str | None = None,
        evidence: str = "",
    ) -> str:
        """Revisa com Gemini e sempre preserva a resposta primária se falhar."""
        gemini_api_key = os.getenv("GEMINI_API_KEY")
        if not gemini_api_key:
            fallback = self.enforce_direct_answer(question, candidate)
            print("Gemini review status: SKIPPED — GEMINI_API_KEY is missing")
            print(f"Primary answer preserved: {fallback}")
            return fallback

        reviewer_model = os.getenv(
            "GAIA_GEMINI_REVIEW_MODEL", DEFAULT_GEMINI_REVIEW_MODEL
        )
        if not reviewer_model.lower().startswith("gemini-"):
            reviewer_model = DEFAULT_GEMINI_REVIEW_MODEL
        review_prompt = f"""
Review this GAIA exact-match candidate. Preserve it unless a correction is
clearly necessary. The final_answer must contain only the requested value,
without explanation, label, Markdown, or citation. Respect requested numeric,
currency, name, list separator/order, quote, or chess notation formats.
Use only the collected evidence when it is present. Do not invent facts and do
not suggest or call another tool.
Task ID: {task_id or "test"}
Question: {question}
Candidate: {candidate}
Collected evidence:
{evidence or "(none; format-check the candidate only)"}
""".strip()

        try:
            response = requests.post(
                (
                    "https://generativelanguage.googleapis.com/v1beta/models/"
                    f"{reviewer_model}:generateContent"
                ),
                headers={
                    "x-goog-api-key": gemini_api_key,
                    "Content-Type": "application/json",
                },
                json={
                    "contents": [
                        {
                            "role": "user",
                            "parts": [{"text": review_prompt}],
                        }
                    ],
                    "generationConfig": {
                        "maxOutputTokens": 512,
                        "thinkingConfig": {"thinkingLevel": "minimal"},
                        "responseMimeType": "application/json",
                        "responseJsonSchema": {
                            "type": "object",
                            "properties": {
                                "final_answer": {
                                    "type": "string",
                                    "description": (
                                        "Only the exact value requested by the "
                                        "question, with no explanation."
                                    ),
                                },
                                "review_note": {
                                    "type": "string",
                                    "description": "Optional very short audit note.",
                                },
                            },
                            "required": ["final_answer"],
                            "additionalProperties": False,
                        },
                    },
                },
                timeout=(WEBPAGE_CONNECT_TIMEOUT, 45),
            )
            response.raise_for_status()
            payload = response.json()
            parts = payload["candidates"][0]["content"]["parts"]
            content = "".join(
                str(part.get("text") or "") for part in parts
            ).strip()
            review_data = json.loads(content)
            final_answer = self.enforce_direct_answer(
                question, str(review_data.get("final_answer") or "")
            )
            if not final_answer:
                raise ValueError("Gemini returned an empty final_answer.")

            note = str(review_data.get("review_note") or "").strip()
            print(f"Candidate answer: {candidate}")
            print(f"Gemini reviewed answer: {final_answer}")
            print(f"Gemini changed answer: {final_answer != candidate}")
            print(f"Gemini review note: {note}")
            return final_answer
        except Exception as exc:
            fallback = self.enforce_direct_answer(question, candidate)
            if not fallback:
                raise RuntimeError(
                    "Gemini review failed and the primary answer was empty. "
                    f"Detail: {compact_error(exc)}"
                ) from exc
            detail = compact_error(exc)
            if "response" in locals() and response is not None:
                detail += f" Response: {response.text[:500]}"
            print(f"Candidate answer: {fallback}")
            print("Gemini review status: FAILED — primary answer preserved")
            print(f"Gemini review error: {detail}")
            return fallback


def empty_results() -> pd.DataFrame:
    return pd.DataFrame(columns=RESULT_COLUMNS)


def runtime_username(profile: gr.OAuthProfile | None = None) -> str:
    if profile and getattr(profile, "username", None):
        return str(profile.username).strip()
    configured = str(os.getenv("HF_USERNAME") or "").strip()
    if configured:
        return configured
    repository = str(os.getenv("GITHUB_REPOSITORY") or "").strip()
    if "/" in repository:
        return repository.split("/", 1)[0].strip()
    return ""


def runtime_has_hf_access(profile: gr.OAuthProfile | None = None) -> bool:
    return bool(profile or str(os.getenv("HF_TOKEN") or "").strip())


def answer_readiness(answer: str) -> tuple[int, str]:
    """Avalia somente se a resposta parece pronta; não verifica o gabarito."""
    text = str(answer or "").strip()
    lowered = text.lower()
    if not text:
        return 0, "resposta vazia"
    if lowered.startswith("error:") or "erro" in lowered[:30]:
        return 0, "erro de execução"
    if lowered in {"not found", "unknown", "não encontrado", "n/a", "none"}:
        return 20, "resposta inconclusiva"
    if "final answer" in lowered:
        return 50, "contém texto proibido"
    if len(text) > 500:
        return 60, "resposta possivelmente longa para exact match"
    return 100, "formato aparentemente pronto"


def readiness_summary(results_table) -> str:
    try:
        if isinstance(results_table, pd.DataFrame):
            dataframe = results_table
        else:
            dataframe = pd.DataFrame(results_table, columns=RESULT_COLUMNS)
        if dataframe.empty:
            return "Índice de prontidão: 0% — nenhuma resposta gerada."

        evaluations = [
            answer_readiness(row.get("Submitted Answer", ""))
            for _, row in dataframe.iterrows()
        ]
        scores = [score for score, _ in evaluations]
        ready = sum(score == 100 for score in scores)
        problems = len(scores) - ready
        average = round(sum(scores) / len(scores))
        return (
            f"Índice de prontidão: {average}% — {ready}/{len(scores)} respostas "
            f"aparentemente prontas; {problems} precisam de revisão. "
            "Este índice avalia erros e formato, não a correção do gabarito. "
            "A nota oficial só existe após o envio."
        )
    except Exception as exc:
        return f"Não foi possível calcular o índice de prontidão: {exc}"


def review_dataframe(questions: list, answers: dict) -> pd.DataFrame:
    rows = []
    for item in questions or []:
        task_id = str(item.get("task_id", "")).strip()
        rows.append(
            {
                "Task ID": task_id,
                "Question": str(item.get("question", "")),
                "Submitted Answer": str((answers or {}).get(task_id, "")),
            }
        )
    return pd.DataFrame(rows, columns=RESULT_COLUMNS)


def progress_summary(questions: list, answers: dict) -> str:
    total = len(questions or [])
    answered = sum(
        bool(str((answers or {}).get(str(item.get("task_id", "")), "")).strip())
        for item in (questions or [])
    )
    return f"Progresso: {answered}/{total} questões respondidas."


def load_evaluation_questions(profile: gr.OAuthProfile | None):
    """Carrega as 20 questões, mas não executa o agente."""
    if not runtime_has_hf_access(profile):
        return (
            "Faça login no Hugging Face ou configure HF_TOKEN no Codespaces.",
            [],
            {},
            gr.update(choices=[], value=None),
            "",
            "",
            "",
            "Progresso: 0/20 questões respondidas.",
            "Índice de prontidão: 0%.",
            empty_results(),
        )

    try:
        questions = fetch_official_questions()
        questions = [
            item
            for item in questions
            if item.get("task_id") and item.get("question") is not None
        ]
        if not questions:
            raise ValueError("A API retornou uma lista vazia.")

        answers = {}
        choices = [
            (
                f"{index + 1:02d}. {str(item['question'])[:90]}",
                str(item["task_id"]),
            )
            for index, item in enumerate(questions)
        ]
        first = questions[0]
        dataframe = review_dataframe(questions, answers)
        return (
            f"{len(questions)} questões carregadas via {QUESTIONS_SOURCE}. "
            "Nenhuma foi executada ainda.",
            questions,
            answers,
            gr.update(choices=choices, value=str(first["task_id"])),
            str(first["question"]),
            str(first["task_id"]),
            "",
            progress_summary(questions, answers),
            readiness_summary(dataframe),
            dataframe,
        )
    except Exception as exc:
        return (
            f"Erro ao carregar questões: {exc}",
            [],
            {},
            gr.update(choices=[], value=None),
            "",
            "",
            "",
            "Progresso: 0/20 questões respondidas.",
            "Índice de prontidão: 0%.",
            empty_results(),
        )


def select_evaluation_question(task_id: str, questions: list, answers: dict):
    for item in questions or []:
        if str(item.get("task_id")) == str(task_id):
            return (
                str(item.get("question", "")),
                str(item.get("task_id", "")),
                str((answers or {}).get(str(task_id), "")),
                "Questão selecionada. Execute o agente ou edite e salve a resposta.",
            )
    return "", "", "", "Questão não encontrada na sessão."


def save_reviewed_answer(
    task_id: str, answer: str, questions: list, answers: dict
):
    task_id = str(task_id or "").strip()
    if not task_id:
        dataframe = review_dataframe(questions, answers)
        return (
            answers or {},
            "Nenhuma questão selecionada.",
            progress_summary(questions, answers),
            readiness_summary(dataframe),
            dataframe,
        )

    updated = dict(answers or {})
    updated[task_id] = str(answer or "").strip()
    dataframe = review_dataframe(questions, updated)
    return (
        updated,
        "Resposta revisada e salva nesta sessão.",
        progress_summary(questions, updated),
        readiness_summary(dataframe),
        dataframe,
    )


def run_current_evaluation_question(
    question: str, task_id: str, questions: list, answers: dict
):
    """Executa somente a questão atualmente selecionada."""
    if not question or not task_id:
        dataframe = review_dataframe(questions, answers)
        return (
            "",
            answers or {},
            "Carregue e selecione uma questão primeiro.",
            progress_summary(questions, answers),
            readiness_summary(dataframe),
            dataframe,
        )

    try:
        answer = BasicAgent()(question, task_id)
        had_previous_answer = bool(
            str((answers or {}).get(str(task_id), "")).strip()
        )
        updated = dict(answers or {})
        updated[str(task_id)] = answer
        dataframe = review_dataframe(questions, updated)
        action_message = (
            "A resposta anterior foi substituída pela nova resposta."
            if had_previous_answer
            else "A primeira resposta desta questão foi salva."
        )
        return (
            answer,
            updated,
            f"Questão executada. {action_message} "
            "Revise o conteúdo antes de avançar.",
            progress_summary(questions, updated),
            readiness_summary(dataframe),
            dataframe,
        )
    except Exception as exc:
        dataframe = review_dataframe(questions, answers)
        return (
            f"ERROR: {exc}",
            answers or {},
            f"Erro ao executar esta questão: {exc}",
            progress_summary(questions, answers),
            readiness_summary(dataframe),
            dataframe,
        )


def run_all_evaluation_questions(
    profile: gr.OAuthProfile | None, questions: list, answers: dict
):
    """Executa todas as questões carregadas e reúne as respostas para revisão."""
    if not runtime_has_hf_access(profile):
        dataframe = review_dataframe(questions, answers)
        return (
            answers or {},
            "Faça login no Hugging Face ou configure HF_TOKEN no Codespaces.",
            progress_summary(questions, answers),
            readiness_summary(dataframe),
            dataframe,
        )

    if not questions:
        dataframe = review_dataframe(questions, answers)
        return (
            answers or {},
            "Primeiro clique em 'Carregar as 20 questões'.",
            progress_summary(questions, answers),
            readiness_summary(dataframe),
            dataframe,
        )

    try:
        agent = BasicAgent()
    except Exception as exc:
        dataframe = review_dataframe(questions, answers)
        return (
            answers or {},
            f"Erro ao inicializar o agente: {exc}",
            progress_summary(questions, answers),
            readiness_summary(dataframe),
            dataframe,
        )

    updated = dict(answers or {})
    failures = 0
    for item in questions:
        task_id = str(item.get("task_id", "")).strip()
        question = str(item.get("question", "")).strip()
        if not task_id or not question:
            continue
        try:
            updated[task_id] = agent(question, task_id)
        except Exception as exc:
            print(
                f"Questão {task_id} falhou e não foi salva: "
                f"{compact_error(exc)}"
            )
            failures += 1

    dataframe = review_dataframe(questions, updated)
    status = (
        f"Execução das {len(questions)} questões concluída. "
        f"Falhas encontradas: {failures}. Revise as respostas antes de enviar."
    )
    return (
        updated,
        status,
        progress_summary(questions, updated),
        readiness_summary(dataframe),
        dataframe,
    )


def fetch_random_question():
    """Busca somente uma questão oficial aleatória, sem executar ou enviar."""
    try:
        item = random.choice(fetch_official_questions())
        task_id = str(item.get("task_id", "")).strip()
        question = str(item.get("question", "")).strip()
        if not task_id or not question:
            raise ValueError("A API retornou uma pergunta em formato inválido.")
        return (
            f"Uma questão GAIA foi carregada via {QUESTIONS_SOURCE}. "
            "Clique em 'Testar esta questão'.",
            question,
            task_id,
            "",
            "",
        )
    except Exception as exc:
        return f"Erro ao buscar questão: {exc}", "", "", "", ""


def test_agent(question: str, task_id: str):
    """Executa uma única pergunta oficial sem enviar a avaliação."""
    if not question or not task_id:
        return "Primeiro carregue uma questão GAIA.", "", ""
    try:
        answer = BasicAgent()(question, task_id)
        score, reason = answer_readiness(answer)
        readiness = (
            f"Índice de prontidão: {score}% ({reason}). "
            "Não é a nota oficial e não compara com o gabarito."
        )
        return "Teste concluído. Nenhum resultado foi enviado.", answer, readiness
    except Exception as exc:
        return f"Erro no teste: {exc}", "", "Índice de prontidão: 0%."


def run_agent_only(profile: gr.OAuthProfile | None):
    """Busca as perguntas e gera uma tabela editável, sem enviar respostas."""
    if not runtime_has_hf_access(profile):
        return (
            "Faça login no Hugging Face ou configure HF_TOKEN no Codespaces.",
            empty_results(),
            "Índice de prontidão: 0%.",
        )

    try:
        agent = BasicAgent()
    except Exception as exc:
        return (
            f"Erro ao inicializar o agente: {exc}",
            empty_results(),
            "Índice de prontidão: 0%.",
        )

    try:
        questions = fetch_official_questions()
        if not questions:
            return (
                "A API retornou uma lista de perguntas vazia.",
                empty_results(),
                "Índice de prontidão: 0%.",
            )
    except Exception as exc:
        return (
            f"Erro ao buscar perguntas: {exc}",
            empty_results(),
            "Índice de prontidão: 0%.",
        )

    results = []
    for item in questions:
        task_id = item.get("task_id")
        question = item.get("question")
        if not task_id or question is None:
            continue

        try:
            answer = agent(question, task_id)
        except Exception as exc:
            print(
                f"Questão {task_id} falhou e foi omitida da avaliação local: "
                f"{compact_error(exc)}"
            )
            continue

        results.append(
            {
                "Task ID": task_id,
                "Question": question,
                "Submitted Answer": answer,
            }
        )

    if not results:
        return (
            "O agente não produziu respostas.",
            empty_results(),
            "Índice de prontidão: 0%.",
        )

    status = (
        f"Execução concluída: {len(results)} respostas geradas. "
        "Revise e, se necessário, edite a coluna 'Submitted Answer'. "
        "Nada foi enviado ainda."
    )
    dataframe = pd.DataFrame(results, columns=RESULT_COLUMNS)
    return status, dataframe, readiness_summary(dataframe)


def normalize_results(results_table) -> list[dict]:
    """Converte a tabela revisada no payload exigido pela API."""
    if results_table is None:
        return []

    if isinstance(results_table, pd.DataFrame):
        dataframe = results_table.copy()
    else:
        dataframe = pd.DataFrame(results_table, columns=RESULT_COLUMNS)

    if dataframe.empty:
        return []

    missing = set(RESULT_COLUMNS) - set(dataframe.columns)
    if missing:
        raise ValueError(
            "A tabela de revisão não contém as colunas esperadas: "
            + ", ".join(sorted(missing))
        )

    answers = []
    for _, row in dataframe.iterrows():
        task_id = str(row["Task ID"]).strip()
        answer = str(row["Submitted Answer"]).strip()
        if not task_id or task_id.lower() == "nan":
            continue
        if not answer or answer.lower() == "nan":
            raise ValueError(f"A tarefa {task_id} está sem resposta.")
        if answer.startswith("ERROR:"):
            raise ValueError(
                f"A tarefa {task_id} ainda contém um erro. "
                "Corrija a resposta antes de enviar."
            )
        answers.append({"task_id": task_id, "submitted_answer": answer})

    return answers


def submit_to_leaderboard(
    profile: gr.OAuthProfile | None, results_table
):
    """Envia exatamente os valores atualmente visíveis na tabela revisada."""
    username = runtime_username(profile)
    if not username:
        return (
            "Envio bloqueado: não foi possível identificar o usuário. "
            "Faça login no Hugging Face ou configure HF_USERNAME."
        )

    try:
        answers = normalize_results(results_table)
    except Exception as exc:
        return f"Envio bloqueado: {exc}"

    if not answers:
        return "Não há respostas para enviar. Execute a avaliação primeiro."

    configured_agent_url = str(os.getenv("AGENT_CODE_URL") or "").strip()
    space_id = str(os.getenv("SPACE_ID") or "").strip()
    github_repository = str(os.getenv("GITHUB_REPOSITORY") or "").strip()
    if configured_agent_url:
        agent_code_url = configured_agent_url
    elif space_id:
        agent_code_url = (
            f"https://huggingface.co/spaces/{space_id}/tree/main"
        )
    elif github_repository:
        agent_code_url = f"https://github.com/{github_repository}"
    else:
        return (
            "Envio bloqueado: não foi possível determinar a URL do agente. "
            "Configure AGENT_CODE_URL ou execute dentro de um repositório "
            "GitHub/Codespaces."
        )

    submission = {
        "username": username,
        "agent_code": agent_code_url,
        "answers": answers,
    }

    try:
        response = requests.post(
            f"{DEFAULT_API_URL}/submit",
            json=submission,
            timeout=90,
        )
        response.raise_for_status()
        result = response.json()
        return (
            "Envio realizado com sucesso!\n"
            f"Usuário: {result.get('username')}\n"
            f"Pontuação: {result.get('score', 'N/A')}% "
            f"({result.get('correct_count', '?')}/"
            f"{result.get('total_attempted', '?')} corretas)\n"
            f"Mensagem: {result.get('message', 'Sem mensagem.')}"
        )
    except requests.exceptions.RequestException as exc:
        detail = ""
        if exc.response is not None:
            detail = f" Resposta da API: {exc.response.text[:500]}"
        return (
            "Falha no envio para a API oficial. As respostas continuam "
            f"salvas na tabela desta sessão. Detalhe: {exc}.{detail}"
        )
    except Exception as exc:
        return f"Erro inesperado no envio: {exc}"


def submit_telegram_answers(
    questions: list[dict], answers: dict[str, str]
) -> str:
    """Converte o estado revisado do Telegram para o mesmo fluxo da interface."""
    return submit_to_leaderboard(
        None,
        review_dataframe(questions, answers),
    )


with gr.Blocks(theme=gr.themes.Soft(), title="GAIA Agent Evaluation") as demo:
    gr.Markdown("# GAIA Agent Evaluation")
    gr.Markdown(
        "Teste o agente isoladamente, gere as respostas oficiais para revisão "
        "e só então faça o envio final."
    )
    gr.LoginButton()

    with gr.Tabs():
        with gr.Tab("1. Testar agente"):
            gr.Markdown(
                "Carregue uma única pergunta aleatória da avaliação GAIA e "
                "teste o agente sem enviar nenhuma resposta."
            )
            random_task_id = gr.State("")
            test_question = gr.Textbox(
                label="Pergunta GAIA sorteada",
                lines=5,
                interactive=False,
            )
            with gr.Row():
                fetch_question_button = gr.Button(
                    "Carregar uma pergunta GAIA", variant="secondary"
                )
                test_button = gr.Button(
                    "Testar esta questão", variant="primary"
                )
            test_status = gr.Textbox(label="Status", interactive=False)
            test_answer = gr.Textbox(
                label="Resposta do agente", lines=5, interactive=False
            )
            test_readiness = gr.Textbox(
                label="Avaliação antes do envio",
                lines=3,
                interactive=False,
            )
            fetch_question_button.click(
                fn=fetch_random_question,
                outputs=[
                    test_status,
                    test_question,
                    random_task_id,
                    test_answer,
                    test_readiness,
                ],
            )
            test_button.click(
                fn=test_agent,
                inputs=[test_question, random_task_id],
                outputs=[test_status, test_answer, test_readiness],
            )

        with gr.Tab("2. Executar e revisar"):
            gr.Markdown(
                "Carregue as 20 questões e execute apenas a questão selecionada. "
                "As respostas ficam reunidas para uma única submissão final."
            )
            evaluation_questions = gr.State([])
            evaluation_answers = gr.State({})
            current_task_id = gr.State("")

            with gr.Row():
                load_questions_button = gr.Button(
                    "Carregar as 20 questões", variant="secondary"
                )
                run_all_button = gr.Button(
                    "Executar as 20 questões", variant="primary"
                )
            question_selector = gr.Dropdown(
                label="Escolha a questão",
                choices=[],
                interactive=True,
            )
            current_question = gr.Textbox(
                label="Questão selecionada",
                lines=6,
                interactive=False,
            )
            current_answer = gr.Textbox(
                label="Resposta da questão selecionada",
                lines=4,
                interactive=True,
                placeholder=(
                    "Execute o agente ou digite/corrija a resposta e clique em salvar."
                ),
            )
            with gr.Row():
                run_current_button = gr.Button(
                    "Responder novamente esta questão", variant="primary"
                )
                save_answer_button = gr.Button("Salvar resposta revisada")

            run_status = gr.Textbox(label="Status", lines=4, interactive=False)
            evaluation_progress = gr.Textbox(
                label="Progresso",
                value="Progresso: 0/20 questões respondidas.",
                interactive=False,
            )
            evaluation_readiness = gr.Textbox(
                label="Avaliação antes do envio",
                lines=4,
                interactive=False,
                value="Índice de prontidão: 0% — nenhuma resposta gerada.",
            )
            results_table = gr.DataFrame(
                headers=RESULT_COLUMNS,
                datatype=["str", "str", "str"],
                value=empty_results(),
                label="Respostas para revisão",
                wrap=True,
                interactive=False,
            )

            load_questions_button.click(
                fn=load_evaluation_questions,
                outputs=[
                    run_status,
                    evaluation_questions,
                    evaluation_answers,
                    question_selector,
                    current_question,
                    current_task_id,
                    current_answer,
                    evaluation_progress,
                    evaluation_readiness,
                    results_table,
                ],
            )
            run_all_button.click(
                fn=run_all_evaluation_questions,
                inputs=[evaluation_questions, evaluation_answers],
                outputs=[
                    evaluation_answers,
                    run_status,
                    evaluation_progress,
                    evaluation_readiness,
                    results_table,
                ],
            )
            question_selector.change(
                fn=select_evaluation_question,
                inputs=[
                    question_selector,
                    evaluation_questions,
                    evaluation_answers,
                ],
                outputs=[
                    current_question,
                    current_task_id,
                    current_answer,
                    run_status,
                ],
            )
            run_current_button.click(
                fn=run_current_evaluation_question,
                inputs=[
                    current_question,
                    current_task_id,
                    evaluation_questions,
                    evaluation_answers,
                ],
                outputs=[
                    current_answer,
                    evaluation_answers,
                    run_status,
                    evaluation_progress,
                    evaluation_readiness,
                    results_table,
                ],
            )
            save_answer_button.click(
                fn=save_reviewed_answer,
                inputs=[
                    current_task_id,
                    current_answer,
                    evaluation_questions,
                    evaluation_answers,
                ],
                outputs=[
                    evaluation_answers,
                    run_status,
                    evaluation_progress,
                    evaluation_readiness,
                    results_table,
                ],
            )

        with gr.Tab("3. Enviar resultado"):
            gr.Markdown(
                "O botão abaixo envia os valores atuais da tabela da aba "
                "anterior. Confira todas as respostas antes de continuar."
            )
            submit_button = gr.Button(
                "Enviar respostas revisadas ao leaderboard",
                variant="primary",
            )
            submit_status = gr.Textbox(
                label="Resultado do envio", lines=6, interactive=False
            )
            submit_button.click(
                fn=submit_to_leaderboard,
                inputs=[results_table],
                outputs=[submit_status],
            )


if __name__ == "__main__":
    if os.getenv("TELEGRAM_BOT_TOKEN"):
        try:
            from telegram_bot import start_telegram_bot

            start_telegram_bot(
                agent_factory=BasicAgent,
                questions_loader=fetch_official_questions,
                scoring_url=DEFAULT_API_URL,
                submission_callback=submit_telegram_answers,
            )
        except Exception as exc:
            print(f"Não foi possível iniciar o bot do Telegram: {exc}")

    demo.launch(
        server_name=os.getenv("GRADIO_SERVER_NAME", "0.0.0.0"),
        server_port=int(
            os.getenv("PORT", os.getenv("GRADIO_SERVER_PORT", "7860"))
        ),
        debug=os.getenv("GRADIO_DEBUG", "true").lower()
        in {"1", "true", "yes", "on"},
        show_error=True,
        share=False,
    )
